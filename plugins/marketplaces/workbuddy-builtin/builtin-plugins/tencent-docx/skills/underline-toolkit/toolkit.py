#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 下划线填空工具库 (Underline Fill-in Toolkit)
=====================================================
仅使用字符下划线（Run 的 underline 属性）实现所有填空场景。

核心设计：
- 所有函数统一支持 value="" 参数
- value="" 或不传 → 空白下划线（模板模式）
- value="xxx" → 填入内容 + 下划线保留（回填模式）

使用例子：（import需要根据具体的项目路径，下面只是一个示例）
    from underline_toolkit import (
        add_underline_run,
        add_normal_run,
        add_paragraph_text,
        setup_a4_page,
    )
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn


# ============================================================
# 核心 API：字符下划线
# ============================================================
def add_normal_run(paragraph, text, font_name="仿宋", font_size=Pt(16), bold=False):
    """
    添加普通文字 Run（无下划线），用于拼接正文。

    Args:
        paragraph: 段落对象
        text: 文字内容
        font_name: 字体
        font_size: 字号
        bold: 是否加粗

    Returns:
        run: 创建的 Run 对象
    """
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    if bold:
        run.font.bold = True
    return run


def add_underline_run(paragraph, blank_spaces=8, font_name="仿宋",
                       font_size=Pt(16), underline_color=None, value=""):
    """
    核心 API：添加带字符下划线的 Run。

    用户在 Word 中只看到一段下划线（空白模式）或带下划线的文字（回填模式）。

    Args:
        paragraph: 段落对象
        blank_spaces: 空白空格数（value 为空时生效，决定下划线长度）
        font_name: 字体
        font_size: 字号
        underline_color: 下划线颜色（可选，6位HEX字符串）
        value: 回填内容。空字符串→空格下划线；非空→内容+下划线保留

    Returns:
        run: 创建的 Run 对象
    """
    # 使用不间断空格(\u00a0 NBSP)作为占位字符
    # NBSP 在 Word/WPS 中被视为"不可裁剪"字符，即使在段落末尾也不会被吞掉
    text = value if value else "\u00a0" * blank_spaces
    run = paragraph.add_run(text)
    # 设置 xml:space="preserve" 以保留空格
    run._element.findall(qn('w:t'))[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.underline = WD_UNDERLINE.SINGLE
    if underline_color:
        rPr = run._element.get_or_add_rPr()
        u = rPr.find(qn('w:u'))
        if u is not None:
            u.set(qn('w:color'), underline_color)

    return run


# ============================================================
# 通用辅助函数
# ============================================================
def add_paragraph_text(doc, text, font_name="仿宋", font_size=Pt(16),
                        bold=False, alignment=None, first_indent=None,
                        space_before=None, space_after=None):
    """
    添加纯文本段落。

    Args:
        doc: Document 对象
        text: 段落文字
        font_name: 字体
        font_size: 字号
        bold: 是否加粗
        alignment: 对齐方式（WD_ALIGN_PARAGRAPH 枚举）
        first_indent: 首行缩进（Pt 值）
        space_before: 段前间距（Pt 值）
        space_after: 段后间距（Pt 值）

    Returns:
        p: 创建的 Paragraph 对象
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    if bold:
        run.font.bold = True
    if alignment is not None:
        p.alignment = alignment
    if first_indent is not None:
        p.paragraph_format.first_line_indent = first_indent
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def setup_a4_page(doc, top_margin=2.54, bottom_margin=2.54,
                   left_margin=3.17, right_margin=3.17):
    """
    设置 A4 页面尺寸和页边距。

    Args:
        doc: Document 对象
        top_margin: 上边距（cm），默认 2.54
        bottom_margin: 下边距（cm），默认 2.54
        left_margin: 左边距（cm），默认 3.17
        right_margin: 右边距（cm），默认 3.17
    """
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)
    section.left_margin = Cm(left_margin)
    section.right_margin = Cm(right_margin)
