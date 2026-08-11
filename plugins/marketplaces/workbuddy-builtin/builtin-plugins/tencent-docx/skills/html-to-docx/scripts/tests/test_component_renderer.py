"""Tests for html_to_docx.components."""
from __future__ import annotations

import warnings

import pytest
from bs4 import BeautifulSoup
from docx import Document

from html_to_docx.components import _registry, register, render_component


def _el(html: str):
    soup = BeautifulSoup(html, "lxml")
    el = soup.find(attrs={"data-component": True})
    if el is None:
        el = soup.find(class_="data-card-grid")
    return el


# ---------------------------------------------------------------------------
# Callout variants
# ---------------------------------------------------------------------------

def test_callout_info():
    doc = Document()
    el = _el('<div data-component="callout" data-variant="info"><div class="callout-content">Info</div></div>')
    assert render_component(el, doc) is True
    assert len(doc.tables) == 1


def test_callout_warning():
    doc = Document()
    el = _el('<div data-component="callout" data-variant="warning"><div class="callout-content">Warning</div></div>')
    assert render_component(el, doc) is True
    assert len(doc.tables) == 1


def test_callout_danger():
    doc = Document()
    el = _el('<div data-component="callout" data-variant="danger"><div class="callout-content">Danger</div></div>')
    assert render_component(el, doc) is True


def test_callout_success():
    doc = Document()
    el = _el('<div data-component="callout" data-variant="success"><div class="callout-content">Success</div></div>')
    assert render_component(el, doc) is True


# ---------------------------------------------------------------------------
# Divider
# ---------------------------------------------------------------------------

def test_divider():
    doc = Document()
    el = _el('<div data-component="divider"></div>')
    render_component(el, doc)
    para = doc.paragraphs[-1]
    assert "w:pBdr" in para._p.xml
    assert "w:bottom" in para._p.xml


# ---------------------------------------------------------------------------
# Section marker
# ---------------------------------------------------------------------------

def test_section_marker():
    doc = Document()
    el = _el('<div data-component="section-marker">Section Title</div>')
    render_component(el, doc)
    para = doc.paragraphs[-1]
    assert "Section Title" in para.text
    assert "w:pBdr" in para._p.xml


# ---------------------------------------------------------------------------
# Data card
# ---------------------------------------------------------------------------

def test_data_card():
    doc = Document()
    el = _el('<div data-component="data-card"><dt>Key</dt><dd>Value</dd></div>')
    render_component(el, doc)
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Key"
    assert doc.tables[0].cell(0, 1).text == "Value"


# ---------------------------------------------------------------------------
# Data card grid
# ---------------------------------------------------------------------------

def test_data_card_grid():
    doc = Document()
    html = (
        '<div class="data-card-grid">'
        '<div data-component="data-card"><dt>A</dt><dd>1</dd></div>'
        '<div data-component="data-card"><dt>B</dt><dd>2</dd></div>'
        '</div>'
    )
    el = BeautifulSoup(html, "lxml").find(class_="data-card-grid")
    # Dispatch via data-component attribute
    el["data-component"] = "data-card-grid"
    render_component(el, doc)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 2


# ---------------------------------------------------------------------------
# Unknown component
# ---------------------------------------------------------------------------

def test_unknown_component_warning():
    doc = Document()
    el = _el('<div data-component="nonexistent">Content</div>')
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        result = render_component(el, doc)
    assert result is False
    assert any("nonexistent" in str(w.message) for w in caught)


# ---------------------------------------------------------------------------
# Dynamic registration
# ---------------------------------------------------------------------------

def test_register_new_component():
    @register("test-custom-xyz")
    def _render(el, doc):
        doc.add_paragraph("custom rendered")

    assert "test-custom-xyz" in _registry
    doc = Document()
    el = _el('<div data-component="test-custom-xyz"></div>')
    render_component(el, doc)
    assert doc.paragraphs[-1].text == "custom rendered"
