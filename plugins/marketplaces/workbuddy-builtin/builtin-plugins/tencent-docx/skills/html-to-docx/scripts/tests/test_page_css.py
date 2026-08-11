"""Tests for html_to_docx.page_css (US2 — CSS @page 页眉页脚 / 命名页继承).

Spec: S-26060126E FR-008~013 / FR-026 / SC-002 / SC-009 / Smoke 3.

TDD: written before implementation (red → green).
"""
from __future__ import annotations

import warnings

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START

from html_to_docx.page_css import (
    MarginBoxPos,
    PageRule,
    StringSetBinding,
    PageCssModel,
    parse_page_css,
    bind_pages_to_sections,
    apply_page_furniture,
)
from html_to_docx.fields import TokenKind
from html_to_docx.section_model import (
    SectionSpec,
    parse_sections,
    mark_section_boundaries,
    apply_sections,
)
from html_to_docx.page_setup import resolve_page_defaults


def _defaults():
    return resolve_page_defaults("<html><body></body></html>", None)


# ---------------------------------------------------------------------------
# parse_page_css
# ---------------------------------------------------------------------------

def test_parse_default_page_rule():
    css = """
    <style>
    @page { @bottom-center { content: counter(page); } }
    </style>
    """
    model = parse_page_css(css)
    assert None in model.rules
    rule = model.rules[None]
    assert MarginBoxPos.BOTTOM_CENTER in rule.margin_boxes
    toks = rule.margin_boxes[MarginBoxPos.BOTTOM_CENTER]
    assert toks[0].kind == TokenKind.COUNTER_PAGE


def test_parse_named_page_rule():
    css = """
    <style>
    @page cover { @bottom-center { content: none; } }
    </style>
    """
    model = parse_page_css(css)
    assert "cover" in model.rules
    rule = model.rules["cover"]
    toks = rule.margin_boxes[MarginBoxPos.BOTTOM_CENTER]
    assert toks[0].kind == TokenKind.NONE


def test_parse_page_binding():
    css = """
    <style>
    section[role="cover"] { page: cover; }
    </style>
    """
    model = parse_page_css(css)
    assert model.page_bindings.get("cover") == "cover"


def test_parse_string_set_stores_raw_selector():
    css = """
    <style>
    h1 { string-set: chapter content(text); }
    </style>
    """
    model = parse_page_css(css)
    assert len(model.string_sets) == 1
    sb = model.string_sets[0]
    assert sb.name == "chapter"
    assert sb.selector == "h1"


def test_parse_all_six_margin_boxes():
    css = """
    <style>
    @page {
      @top-left { content: "TL"; }
      @top-center { content: "TC"; }
      @top-right { content: "TR"; }
      @bottom-left { content: "BL"; }
      @bottom-center { content: "BC"; }
      @bottom-right { content: "BR"; }
    }
    </style>
    """
    model = parse_page_css(css)
    rule = model.rules[None]
    assert set(rule.margin_boxes.keys()) == set(MarginBoxPos)


def test_parse_unsupported_syntax_warns_no_abort():
    """FR-013 / SC-009: @page :first pseudo-class -> ignored + warning."""
    css = """
    <style>
    @page :first { @top-center { content: "x"; } }
    @page { @bottom-center { content: counter(page); } }
    </style>
    """
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        model = parse_page_css(css)
    # The valid default @page is still parsed.
    assert None in model.rules
    # A warning was emitted for the unsupported :first.
    assert any("first" in str(w.message).lower()
               or "unsupported" in str(w.message).lower() for w in caught)


def test_parse_empty_string_content_warns():
    """FR-026: content: "" is not a legal value -> warning."""
    css = """
    <style>
    @page { @bottom-center { content: ""; } }
    </style>
    """
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        parse_page_css(css)
    assert any("empty" in str(w.message).lower() for w in caught)


# ---------------------------------------------------------------------------
# bind_pages_to_sections
# ---------------------------------------------------------------------------

def test_bind_pages_to_sections_sets_page_name():
    model = PageCssModel(
        rules={None: PageRule(), "cover": PageRule(name="cover")},
        page_bindings={"cover": "cover"},
    )
    specs = [SectionSpec(index=0, role="cover"), SectionSpec(index=1, role="body")]
    bind_pages_to_sections(model, specs)
    assert specs[0].page_name == "cover"
    assert specs[1].page_name is None


def test_bind_missing_named_page_fallback_warns():
    model = PageCssModel(
        rules={None: PageRule()},
        page_bindings={"cover": "nonexistent"},
    )
    specs = [SectionSpec(index=0, role="cover")]
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        bind_pages_to_sections(model, specs)
    assert specs[0].page_name is None
    assert any("nonexistent" in str(w.message).lower()
               or "fallback" in str(w.message).lower() for w in caught)


# ---------------------------------------------------------------------------
# apply_page_furniture — end-to-end (model C)
# ---------------------------------------------------------------------------

def _build_model_c_doc():
    """Cover (no furniture) + body (page number footer)."""
    html = (
        "<html><head><style>"
        "@page { @bottom-center { content: counter(page); } }"
        "@page cover { @bottom-center { content: none; } }"
        "section[role=\"cover\"] { page: cover; }"
        "</style></head><body>"
        "<section role='cover'><h1>Cover</h1></section>"
        "<section role='body'><h1>Body</h1><p>x</p></section>"
        "</body></html>"
    )
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    specs = parse_sections(soup)
    model = parse_page_css(html)
    bind_pages_to_sections(model, specs)
    marked = mark_section_boundaries(soup, specs)
    from html4docx import HtmlToDocx
    doc = Document()
    HtmlToDocx().add_html_to_document(marked, doc)
    apply_sections(doc, specs, _defaults())
    apply_page_furniture(doc, model, specs)
    return doc, specs, model


def test_model_c_cover_no_footer_body_has_pagenum():
    """AC1 / SC-002: cover section has no footer content; body has page field."""
    doc, specs, model = _build_model_c_doc()
    cover_footer = doc.sections[0].footer
    body_footer = doc.sections[1].footer
    # Cover footer empty (no furniture).
    assert cover_footer.paragraphs[0].text.strip() == ""
    # Body footer has a PAGE field.
    body_xml = body_footer._element.xml
    assert "PAGE" in body_xml or "fldSimple" in body_xml


def test_model_c_cover_footer_not_linked():
    """Cover (all-none) section does not link/create a footer part."""
    doc, specs, model = _build_model_c_doc()
    # Cover should have no real footer content.
    cover_footer = doc.sections[0].footer
    assert cover_footer.is_linked_to_previous or \
        cover_footer.paragraphs[0].text.strip() == ""


def test_string_ref_top_box_writes_styleref():
    """AC2: @top-right { content: string(chapter) } -> header STYLEREF field."""
    html = (
        "<html><head><style>"
        "@page { @top-right { content: string(chapter); } }"
        "h1 { string-set: chapter content(text); }"
        "</style></head><body>"
        "<section role='body'><h1>Chapter</h1><p>x</p></section>"
        "</body></html>"
    )
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    specs = parse_sections(soup)
    model = parse_page_css(html)
    bind_pages_to_sections(model, specs)
    marked = mark_section_boundaries(soup, specs)
    from html4docx import HtmlToDocx
    doc = Document()
    HtmlToDocx().add_html_to_document(marked, doc)
    apply_sections(doc, specs, _defaults())
    apply_page_furniture(doc, model, specs)
    header_xml = doc.sections[0].header._element.xml
    assert "STYLEREF" in header_xml


def test_section_role_cover_missing_named_page_fallback():
    """Edge: section[role=cover]{page:cover} but no @page cover -> fallback + warning."""
    html = (
        "<html><head><style>"
        "@page { @bottom-center { content: counter(page); } }"
        "section[role=\"cover\"] { page: cover; }"
        "</style></head><body>"
        "<section role='cover'><h1>C</h1></section>"
        "</body></html>"
    )
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    specs = parse_sections(soup)
    model = parse_page_css(html)
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        bind_pages_to_sections(model, specs)
    assert specs[0].page_name is None
