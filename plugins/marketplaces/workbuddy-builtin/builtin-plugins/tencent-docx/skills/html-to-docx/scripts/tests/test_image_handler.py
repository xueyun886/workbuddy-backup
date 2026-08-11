"""Tests for html_to_docx.image_handler."""
from __future__ import annotations

import base64
import io
import os

import pytest
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm

from html_to_docx.image_handler import embed_image

# Minimal valid 1×1 PNG (base64)
_PNG_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
_PNG_DATA_URI = f"data:image/png;base64,{_PNG_B64}"


def _img_tag(src: str, width: str | None = None) -> object:
    attrs = f'src="{src}"'
    if width:
        attrs += f' width="{width}"'
    soup = BeautifulSoup(f"<img {attrs}>", "lxml")
    return soup.find("img")


def test_base64_embed():
    """Base64 data URI is decoded and embedded as a picture."""
    doc = Document()
    embed_image(_img_tag(_PNG_DATA_URI), doc)
    # An inline picture adds a paragraph with a drawing
    xml = doc.element.xml
    assert "w:drawing" in xml or "pic:pic" in xml or len(doc.inline_shapes) > 0


def test_local_path_missing():
    """Missing local file produces fallback paragraph."""
    doc = Document()
    embed_image(_img_tag("/nonexistent/does_not_exist.png"), doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("[图片无法加载]" in t for t in texts)


def test_remote_url_timeout(monkeypatch):
    """Remote URL that raises ConnectTimeout produces fallback paragraph with URL."""
    import httpx
    from html_to_docx import image_handler

    def _mock_get(url, **kwargs):
        raise httpx.ConnectTimeout("timeout", request=None)

    monkeypatch.setattr(httpx, "get", _mock_get)

    doc = Document()
    url = "https://example.com/image.png"
    embed_image(_img_tag(url), doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("example.com" in t for t in texts)


def test_oversized_scale():
    """Image with width > max_width_cm is scaled down."""
    doc = Document()
    # Request a 3000px-wide image (≈ 79 cm), max is 16 cm
    embed_image(_img_tag(_PNG_DATA_URI, width="3000"), doc, max_width_cm=16.0)
    # The picture should exist; verify no exception was raised and a picture was added
    xml = doc.element.xml
    assert "w:drawing" in xml or "pic:pic" in xml or len(doc.inline_shapes) > 0
