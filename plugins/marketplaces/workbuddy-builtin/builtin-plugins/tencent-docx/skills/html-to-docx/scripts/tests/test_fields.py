"""Tests for html_to_docx.fields (US3 — 动态域 PAGE / NUMPAGES / STYLEREF).

Spec: S-26060126E FR-014~017 / D-SP-08 / D-SP-10 / D-R3 / SC-003~005.

TDD: written to lock the field-writing contract (red → green).
"""
from __future__ import annotations

import warnings

import pytest
from docx import Document

from html_to_docx.fields import (
    TokenKind,
    PageContentToken,
    tokenize_content,
    resolve_styleref_style_id,
    write_field,
    enable_update_fields,
)
from html_to_docx.page_css import StringSetBinding


# ---------------------------------------------------------------------------
# tokenize_content
# ---------------------------------------------------------------------------

def test_tokenize_counter_literal_combo():
    """AC1: counter(page) " / " counter(pages) -> ordered token list."""
    toks = tokenize_content('counter(page) " / " counter(pages)')
    kinds = [t.kind for t in toks]
    assert kinds == [
        TokenKind.COUNTER_PAGE,
        TokenKind.LITERAL,
        TokenKind.COUNTER_PAGES,
    ]
    assert toks[1].value == " / "


def test_tokenize_string_ref():
    toks = tokenize_content("string(chapter)")
    assert len(toks) == 1
    assert toks[0].kind == TokenKind.STRING_REF
    assert toks[0].value == "chapter"


def test_tokenize_none():
    toks = tokenize_content("none")
    assert len(toks) == 1
    assert toks[0].kind == TokenKind.NONE


def test_tokenize_single_quoted_literal():
    toks = tokenize_content("'hello'")
    assert toks[0].kind == TokenKind.LITERAL
    assert toks[0].value == "hello"


def test_tokenize_unrecognized_fragment_warns_no_abort():
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        toks = tokenize_content('@@@ counter(page)')
    # The valid counter(page) still parses.
    assert any(t.kind == TokenKind.COUNTER_PAGE for t in toks)
    assert any("unrecognized" in str(w.message).lower() for w in caught)


# ---------------------------------------------------------------------------
# resolve_styleref_style_id (FR-016~017 / D-SP-08 / SC-004)
# ---------------------------------------------------------------------------

def test_resolve_styleref_h1_to_heading1():
    """AC2: string-set on h1 -> styleId 'Heading1' (not display name)."""
    sets = [StringSetBinding(name="chapter", selector="h1")]
    assert resolve_styleref_style_id("chapter", sets) == "Heading1"


def test_resolve_styleref_h2_to_heading2():
    """AC3: string-set on h2 -> 'Heading2' (level not hard-coded), SC-004."""
    sets = [StringSetBinding(name="sec", selector="h2")]
    assert resolve_styleref_style_id("sec", sets) == "Heading2"


def test_resolve_styleref_unbound_name_returns_none_warns():
    """Edge: name never bound -> None + warning (field still writable empty)."""
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        result = resolve_styleref_style_id("ghost", [])
    assert result is None
    assert any("ghost" in str(w.message).lower() for w in caught)


@pytest.mark.parametrize("selector", ["h1.x", ".x", "#x", "div"])
def test_resolve_styleref_non_bare_heading_returns_none(selector):
    """Edge: compound / class / id / non-heading selector -> None + warning."""
    sets = [StringSetBinding(name="c", selector=selector)]
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        result = resolve_styleref_style_id("c", sets)
    assert result is None
    assert len(caught) >= 1


# ---------------------------------------------------------------------------
# write_field — OOXML field structures (D-R3)
# ---------------------------------------------------------------------------

def test_write_field_page_emits_fldsimple_page():
    """AC1: COUNTER_PAGE -> <w:fldSimple w:instr="PAGE">."""
    doc = Document()
    para = doc.add_paragraph()
    write_field(para, [PageContentToken(TokenKind.COUNTER_PAGE)], None)
    xml = para._p.xml
    assert "fldSimple" in xml
    assert "PAGE" in xml


def test_write_field_numpages_emits_numpages_native():
    """AC5: COUNTER_PAGES -> NUMPAGES native total (no -1, D-SP-10)."""
    doc = Document()
    para = doc.add_paragraph()
    write_field(para, [PageContentToken(TokenKind.COUNTER_PAGES)], None)
    xml = para._p.xml
    assert "NUMPAGES" in xml
    # Native count: no arithmetic subtraction baked into the field.
    assert "- 1" not in xml and "-1" not in xml


def test_write_field_string_ref_emits_styleref_styleid():
    """AC2: STRING_REF resolves to STYLEREF <styleId> three-part field."""
    doc = Document()
    para = doc.add_paragraph()
    sets = [StringSetBinding(name="chapter", selector="h1")]

    def resolver(name):
        return resolve_styleref_style_id(name, sets)

    write_field(para, [PageContentToken(TokenKind.STRING_REF, "chapter")], resolver)
    xml = para._p.xml
    assert "STYLEREF" in xml
    assert "Heading1" in xml
    # Three-part complex field: fldChar begin + instrText + fldChar end.
    assert 'fldCharType="begin"' in xml
    assert 'fldCharType="end"' in xml
    assert "instrText" in xml


def test_write_field_string_ref_unresolved_writes_empty_styleref():
    """Edge: unresolved STYLEREF still writes a (empty) field, never aborts."""
    doc = Document()
    para = doc.add_paragraph()
    write_field(para, [PageContentToken(TokenKind.STRING_REF, "ghost")], lambda n: None)
    xml = para._p.xml
    assert "STYLEREF" in xml


def test_write_field_literal_emits_plain_run():
    doc = Document()
    para = doc.add_paragraph()
    write_field(para, [PageContentToken(TokenKind.LITERAL, "Page ")], None)
    assert para.text == "Page "


def test_write_field_none_emits_nothing():
    doc = Document()
    para = doc.add_paragraph()
    write_field(para, [PageContentToken(TokenKind.NONE)], None)
    assert para.text == ""
    assert "fldSimple" not in para._p.xml


def test_write_field_reopen_valid_ooxml():
    """Invariant: written field structure survives a python-docx round-trip."""
    import io

    doc = Document()
    para = doc.add_paragraph()
    write_field(
        para,
        [
            PageContentToken(TokenKind.COUNTER_PAGE),
            PageContentToken(TokenKind.LITERAL, " / "),
            PageContentToken(TokenKind.COUNTER_PAGES),
        ],
        None,
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    reopened = Document(buf)
    assert reopened is not None


# ---------------------------------------------------------------------------
# enable_update_fields (D-R3 / SC-003)
# ---------------------------------------------------------------------------

def test_enable_update_fields_sets_settings_flag():
    doc = Document()
    enable_update_fields(doc)
    from docx.oxml.ns import qn
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    assert el is not None
    assert el.get(qn("w:val")) == "true"


def test_enable_update_fields_idempotent():
    doc = Document()
    enable_update_fields(doc)
    enable_update_fields(doc)
    from docx.oxml.ns import qn
    settings = doc.settings.element
    els = settings.findall(qn("w:updateFields"))
    assert len(els) == 1
