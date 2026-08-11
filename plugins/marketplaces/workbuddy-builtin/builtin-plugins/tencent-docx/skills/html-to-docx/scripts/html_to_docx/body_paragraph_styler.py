"""body_paragraph_styler.py — Phase 2d (post): apply body paragraph styles.

命门集成 (T020b / FR-001~003 端到端)：

US2 (em) + US1 (border/shd) 仅扩展了 ``style_mapper.apply_paragraph_styles``，
但正文普通 ``<p>/<div>/<h1>~<h6>/<li>`` 由第三方 ``html4docx`` 渲染，**从不调用**
``apply_paragraph_styles`` → 段落边框 / 底纹 / em 段距在正文上端到端丢失。

本模块镜像 ``table_style_applier.apply_table_styles`` 的「html4docx 后处理」范式：
  1. 解析 clean HTML（style_injector 已把 class 样式内联到元素 ``style``），
     收集**叶子级块级元素**（不含其他块级子元素，对应 docx 中一个独立段落）。
  2. 仅保留带「段落级视觉样式」（border / background / margin / text-indent /
     line-height / font-size）的元素作为候选。
  3. 顺序双指针把候选元素匹配到 ``document.paragraphs``（python-docx 的
     ``document.paragraphs`` 天然只含 body 顶层段落，**不含表格单元格段落**，
     故与 ``table_style_applier`` 天然隔离、不重复处理表格内段落）。
  4. 对匹配到的段落调用 ``apply_paragraph_styles``（**复用** US1/US2 实现，不重造）。

落点：``converter.py`` Phase 2d（``apply_table_styles``）之后挂载。
"""
from __future__ import annotations

import re
from typing import TYPE_CHECKING

from bs4 import BeautifulSoup, Tag

from .style_mapper import apply_paragraph_styles

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph


# 块级元素：每个叶子级实例对应 docx 中一个独立段落。
_BLOCK_TAGS = ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li")

# 判定「含块级/表格子元素」时一并纳入的容器标签（这些子节点会另起段落/表格）。
_NESTED_BLOCK_TAGS = (
    "p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li",
    "table", "ul", "ol", "section", "main", "article", "header", "footer",
)

# 仅当元素带这些「段落级视觉样式」属性时才需要 post-pass（避免无谓处理）。
# 与 style_mapper.apply_paragraph_styles 关心的属性集合对齐。
_RELEVANT_PREFIXES = ("border", "background", "margin", "text-indent", "line-height", "font-size")

# converter.py 用于标记特殊区域占位段落的前缀（跳过，避免误匹配）。
_PLACEHOLDER_PREFIX = "\u200b__SPECIAL_REGION__:"

# 不可见的 section boundary marker（section_model 注入），归一化时一并剔除。
_ZERO_WIDTH = "\u200b\u200c\u200d\ufeff"


def _normalize_text(text: str) -> str:
    """归一化文本用于顺序匹配：去零宽字符与所有空白。"""
    if not text:
        return ""
    for ch in _ZERO_WIDTH:
        text = text.replace(ch, "")
    return re.sub(r"\s+", "", text)


def _has_relevant_style(props: dict[str, str]) -> bool:
    """元素是否带需要 post-pass 处理的段落级视觉样式。"""
    for key in props:
        if any(key == p or key.startswith(p) for p in _RELEVANT_PREFIXES):
            return True
    return False


def _parse_style(style_attr: str) -> dict[str, str]:
    props: dict[str, str] = {}
    for decl in style_attr.split(";"):
        decl = decl.strip()
        if ":" in decl:
            k, _, v = decl.partition(":")
            props[k.strip().lower()] = v.strip()
    return props


def _is_inside_table(el: Tag) -> bool:
    """元素是否位于 <table> 内（表格单元格段落由 table_style_applier 负责）。"""
    for parent in el.parents:
        if isinstance(parent, Tag) and parent.name == "table":
            return True
    return False


def _is_leaf_block(el: Tag) -> bool:
    """叶子级块元素：不含会另起段落/表格的块级子元素 → 对应一个独立 docx 段落。"""
    return el.find(_NESTED_BLOCK_TAGS) is None


def _collect_styled_blocks(html: str) -> list[tuple[str, str]]:
    """按文档顺序收集 (normalized_text, style_attr) 候选。

    仅保留：叶子级块元素 + 不在表格内 + 带段落级视觉样式 + 有可见文本。
    """
    soup = BeautifulSoup(html, "lxml")
    out: list[tuple[str, str]] = []
    for el in soup.find_all(_BLOCK_TAGS):
        style_attr = el.get("style", "")
        if not style_attr:
            continue
        props = _parse_style(style_attr)
        if not _has_relevant_style(props):
            continue
        if not _is_leaf_block(el):
            continue
        if _is_inside_table(el):
            continue
        norm = _normalize_text(el.get_text())
        if not norm:
            continue  # 空文本段落无法稳定匹配，跳过
        out.append((norm, style_attr))
    return out


def _iter_body_paragraphs(document: DocxDocument):
    """遍历 body 顶层段落，跳过占位段落（component/toc placeholder）。

    python-docx 的 ``document.paragraphs`` 不含表格单元格内段落，天然与
    ``table_style_applier`` 隔离。
    """
    for para in document.paragraphs:
        text = para.text
        if text.startswith(_PLACEHOLDER_PREFIX):
            continue
        yield para


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def apply_body_paragraph_styles(document: DocxDocument, html: str) -> None:
    """把 clean HTML 正文块级元素的段落级样式施加到对应 docx 段落（FR-001~003）。

    顺序双指针匹配：对每个候选块元素，从当前 docx 段落指针起向后找首个
    归一化文本匹配的段落，调用 ``apply_paragraph_styles`` 后推进指针。

    Args:
        document: html4docx 渲染后的 docx 文档。
        html: clean HTML（style_injector 内联后、placeholder 替换前的快照）。
    """
    blocks = _collect_styled_blocks(html)
    if not blocks:
        return

    paragraphs: list[Paragraph] = list(_iter_body_paragraphs(document))
    if not paragraphs:
        return

    para_idx = 0
    for norm_text, style_attr in blocks:
        # 从当前指针向后找首个文本匹配的段落（顺序匹配，避免错位）。
        match_idx = -1
        for i in range(para_idx, len(paragraphs)):
            ptext = _normalize_text(paragraphs[i].text)
            if not ptext:
                continue
            if ptext == norm_text or norm_text in ptext or ptext in norm_text:
                match_idx = i
                break
        if match_idx == -1:
            # 未匹配（如内容被组件/占位吃掉）→ 跳过该块，不推进指针。
            continue
        apply_paragraph_styles(paragraphs[match_idx], style_attr)
        para_idx = match_idx + 1
