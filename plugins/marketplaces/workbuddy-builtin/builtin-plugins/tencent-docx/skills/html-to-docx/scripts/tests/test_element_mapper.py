"""Tests for html_to_docx.element_mapper."""
from __future__ import annotations

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from html_to_docx.element_mapper import (
    add_blockquote,
    add_caption,
    add_horizontal_rule,
    apply_subscript,
    apply_superscript,
    merge_cells,
)


def test_caption():
    """add_caption produces a centred paragraph with muted small font."""
    doc = Document()
    add_caption(doc, "Table 1: Test Caption")
    para = doc.paragraphs[-1]
    assert "Test Caption" in para.text
    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    assert run.font.size == Pt(9)


def test_superscript_subscript():
    """apply_superscript/subscript set the correct font properties."""
    doc = Document()
    para = doc.add_paragraph("base")
    run_sup = para.add_run("sup")
    apply_superscript(run_sup)
    assert run_sup.font.superscript is True

    run_sub = para.add_run("sub")
    apply_subscript(run_sub)
    assert run_sub.font.subscript is True


def test_colspan_rowspan():
    """merge_cells merges a 2×2 block without raising."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    merge_cells(table, 0, 0, rowspan=2, colspan=2)
    # Merging is destructive; just verify no exception and row count intact
    assert len(table.rows) == 3


def test_hr():
    """add_horizontal_rule adds a paragraph with a bottom border."""
    doc = Document()
    add_horizontal_rule(doc)
    para = doc.paragraphs[-1]
    xml = para._p.xml
    assert "w:pBdr" in xml
    assert "w:bottom" in xml


def test_blockquote():
    """add_blockquote adds a left-indented paragraph with a left border."""
    doc = Document()
    add_blockquote(doc, "Quoted text here.")
    para = doc.paragraphs[-1]
    assert "Quoted text" in para.text
    assert abs(para.paragraph_format.left_indent - Cm(1)) < 100  # ±100 EMU tolerance
    xml = para._p.xml
    assert "w:left" in xml
