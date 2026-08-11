"""test_body_paragraph_integration.py — T020b: 正文段落集成 post-pass（命门）。

US2 em + US1 border/shd 仅改了 ``apply_paragraph_styles``，但正文普通
``<p>/<div>`` 走 html4docx 从不调用它 → 端到端仍丢边框/底纹/段距。

本测试走**完整 Converter 流程**（HTML → docx 文件），断言正文段落（非组件、
非表格单元格）真实带上：
  - pPr/pBdr/left（w:val/w:sz/w:color 匹配）  ← FR-002 端到端
  - pPr/shd@w:fill                              ← FR-003 端到端
  - space_after == Pt(11)（0.5em @ 22pt）        ← FR-001 端到端

并断言无 border/bg/margin 的普通段落**不受影响**（无 pBdr/shd，回归安全）。
"""
from __future__ import annotations

import os
import tempfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from html_to_docx.converter import Converter
from html_to_docx.types import ConvertOptions


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _convert(html: str) -> Document:
    out = tempfile.mktemp(suffix=".docx")
    result = Converter(ConvertOptions(output_path=out)).convert(html)
    assert result.success, result.error
    try:
        return Document(out)
    finally:
        if os.path.exists(out):
            os.unlink(out)


def _find_para(doc: Document, text_contains: str):
    for p in doc.paragraphs:
        if text_contains in p.text:
            return p
    return None


def _pBdr(para):
    pPr = para._p.find(qn("w:pPr"))
    return None if pPr is None else pPr.find(qn("w:pBdr"))


def _shd(para):
    pPr = para._p.find(qn("w:pPr"))
    return None if pPr is None else pPr.find(qn("w:shd"))


# ---------------------------------------------------------------------------
# 命门：带 inline style 的正文段落 端到端保真
# ---------------------------------------------------------------------------

_HTML_STYLED_BODY = """
<html><head></head><body>
<p style="border-left:3px solid #1a4d8f;background:#eef4fb;margin-bottom:0.5em;font-size:22pt">核心观点正文段落</p>
<p>普通段落无任何视觉样式</p>
</body></html>
"""


def test_body_paragraph_border_applied_end_to_end():
    doc = _convert(_HTML_STYLED_BODY)
    para = _find_para(doc, "核心观点正文段落")
    assert para is not None, "样式正文段落未在 docx 中找到"
    pBdr = _pBdr(para)
    assert pBdr is not None, "正文段落 border-left 端到端丢失（命门未集成）"
    left = pBdr.find(qn("w:left"))
    assert left is not None
    assert left.get(qn("w:val")) == "single"
    # 3px → 2.25pt → *8 = 18
    assert left.get(qn("w:sz")) == "18"
    assert left.get(qn("w:color")) == "1A4D8F"


def test_body_paragraph_shading_applied_end_to_end():
    doc = _convert(_HTML_STYLED_BODY)
    para = _find_para(doc, "核心观点正文段落")
    assert para is not None
    shd = _shd(para)
    assert shd is not None, "正文段落 background 端到端丢失（命门未集成）"
    assert shd.get(qn("w:fill")) == "EEF4FB"


def test_body_paragraph_em_space_after_end_to_end():
    doc = _convert(_HTML_STYLED_BODY)
    para = _find_para(doc, "核心观点正文段落")
    assert para is not None
    # margin-bottom:0.5em @ font-size:22pt → 11pt（em 按元素字号换算）
    assert para.paragraph_format.space_after == Pt(11)


def test_plain_body_paragraph_untouched():
    """无 border/bg/margin 的普通段落不应被 post-pass 误加 pBdr/shd。"""
    doc = _convert(_HTML_STYLED_BODY)
    para = _find_para(doc, "普通段落无任何视觉样式")
    assert para is not None
    assert _pBdr(para) is None, "普通段落被误加边框"
    assert _shd(para) is None, "普通段落被误加底纹"


# ---------------------------------------------------------------------------
# class 注入路径（贴近真实模板：样式由 <style> class 注入而非内联）
# ---------------------------------------------------------------------------

_HTML_CLASS_STYLED = """
<html><head><style>
.callout-box { border-left:3px solid #1a4d8f; background-color:#eef4fb; margin-bottom:0.5em; font-size:22pt; }
.section-header { background:#1a237e; color:#ffffff; font-size:13pt; }
</style></head><body>
<div class="section-header">章节标题文本</div>
<p class="callout-box">通过 class 注入样式的正文段落</p>
<p>另一个普通段落</p>
</body></html>
"""


def test_class_injected_paragraph_border_shd_end_to_end():
    doc = _convert(_HTML_CLASS_STYLED)
    para = _find_para(doc, "通过 class 注入样式的正文段落")
    assert para is not None
    pBdr = _pBdr(para)
    assert pBdr is not None, "class 注入的 border 端到端丢失"
    assert pBdr.find(qn("w:left")) is not None
    shd = _shd(para)
    assert shd is not None, "class 注入的 background 端到端丢失"
    assert shd.get(qn("w:fill")) == "EEF4FB"


def test_class_injected_section_header_shd_end_to_end():
    """章节标题 div（深底白字）的背景色端到端不丢失（核心诉求②正文上落地）。"""
    doc = _convert(_HTML_CLASS_STYLED)
    para = _find_para(doc, "章节标题文本")
    assert para is not None
    shd = _shd(para)
    assert shd is not None, "section-header 深底背景端到端丢失"
    assert shd.get(qn("w:fill")) == "1A237E"


# ---------------------------------------------------------------------------
# 协同：表格单元格内段落不被 body post-pass 重复处理
# ---------------------------------------------------------------------------

_HTML_WITH_TABLE = """
<html><head></head><body>
<p style="border-left:2px solid #1a4d8f">表外正文段落</p>
<table>
  <tr><td style="background-color:#dddddd">单元格文本内容</td></tr>
</table>
</body></html>
"""


_HTML_ABSTRACT_CARD = """
<html><head><style>
.abstract-card { width:100%; }
.abstract-card td { background-color:#eef4fb; border-left:3px solid #1a4d8f; padding:0.6em 1em; }
</style></head><body>
<table class="abstract-card"><tr><td>
<p><strong>核心观点</strong></p>
<p>第一段富文本内容。</p>
<p>第二段富文本内容。</p>
</td></tr></table>
</body></html>
"""


def test_abstract_card_cell_padding_tcmar():
    """裸 table 卡片单元格 padding → w:tcMar 内边距（D-TF-02 缺口补遗）。"""
    doc = _convert(_HTML_ABSTRACT_CARD)
    assert len(doc.tables) == 1
    cell = doc.tables[0].rows[0].cells[0]
    tcPr = cell._tc.find(qn("w:tcPr"))
    assert tcPr is not None
    # shd + 富文本多段落（回归基线已有）
    assert tcPr.find(qn("w:shd")) is not None
    # 多段落富文本保留（含标题段）
    texts = [p.text for p in cell.paragraphs if p.text.strip()]
    assert "核心观点" in texts
    assert any("第一段" in t for t in texts)
    assert any("第二段" in t for t in texts)
    # 命门补遗：padding → tcMar 必须落地
    tcMar = tcPr.find(qn("w:tcMar"))
    assert tcMar is not None, "单元格 padding 未映射为 w:tcMar（D-TF-02 缺口）"
    # 上下左右四边都应写入（padding:0.6em 1em → 上下0.6em / 左右1em）
    for side in ("top", "bottom", "left", "right"):
        node = tcMar.find(qn(f"w:{side}"))
        assert node is not None, f"tcMar 缺 {side}"
        # w 单位 = 1/20 pt（twips），应 > 0
        assert int(node.get(qn("w:w"))) > 0


def test_table_cell_paragraph_not_double_processed():
    """表格单元格内段落由 table_style_applier 负责，body post-pass 不应触碰。

    document.paragraphs 天然只含 body 顶层段落（不含表格单元格段落），
    因此 body post-pass 解析 clean_html 块级元素时也必须排除 table 内的块级元素。
    """
    doc = _convert(_HTML_WITH_TABLE)
    # 表外段落应有 border
    p_outside = _find_para(doc, "表外正文段落")
    assert p_outside is not None
    assert _pBdr(p_outside) is not None
    # 表格单元格底纹由 table_style_applier 施加（验证未被破坏）
    assert len(doc.tables) == 1
    cell = doc.tables[0].rows[0].cells[0]
    tcPr = cell._tc.find(qn("w:tcPr"))
    assert tcPr is not None
    cell_shd = tcPr.find(qn("w:shd"))
    assert cell_shd is not None
    assert cell_shd.get(qn("w:fill")) == "DDDDDD"


# ===========================================================================
# T030: Quickstart 端到端最小示例
# ===========================================================================

_HTML_QUICKSTART = """<!doctype html>
<html><head><style>
:root { --typography-fontSize-body: 9pt; }
body { font-size: 9pt; line-height: 1.7; }
.abstract-card td { background-color:#eef4fb; border-left:3px solid #1a4d8f; padding:0.6em 1em; }
</style></head>
<body>
  <!-- US2 em：0.5em@22pt 应 = 11pt -->
  <p style="font-size:22pt; margin-bottom:0.5em">大标题em段后距</p>

  <!-- US1 段落 border-left -->
  <p style="border-left:3px solid #1a4d8f; padding-left:8px">带左边框强调段</p>

  <!-- US1 连续相同 border-left -->
  <p style="border-left:3px solid #1a4d8f">连续段一</p>
  <p style="border-left:3px solid #1a4d8f">连续段二</p>

  <!-- US1 段落 background -->
  <p style="background-color:#eef4fb">浅蓝底纹段</p>

  <!-- US1-4 多段落底纹卡片 = 裸单列 table -->
  <table class="abstract-card"><tr><td>
    <p><strong>核心观点</strong></p>
    <p>第一段论述文本</p>
    <p>第二段论述文本</p>
  </td></tr></table>

  <!-- US3 callout 富文本 -->
  <div data-component="callout" data-variant="info">
    <div class="callout-content">
      <p><strong>提示标题</strong></p>
      <p>正文段一</p>
      <p>正文段二</p>
    </div>
  </div>
</body></html>
"""


class TestQuickstartEndToEnd:
    """T030: quickstart.md 端到端最小示例验证."""

    def test_em_space_after(self):
        """US2: font-size:22pt + margin-bottom:0.5em → space_after == Pt(11)."""
        doc = _convert(_HTML_QUICKSTART)
        para = _find_para(doc, "大标题em段后距")
        assert para is not None, "em 段落未找到"
        assert para.paragraph_format.space_after == Pt(11)

    def test_paragraph_border_left(self):
        """US1: border-left:3px solid #1a4d8f → pBdr/left 色值匹配."""
        doc = _convert(_HTML_QUICKSTART)
        para = _find_para(doc, "带左边框强调段")
        assert para is not None
        pBdr = _pBdr(para)
        assert pBdr is not None, "border-left 丢失"
        left = pBdr.find(qn("w:left"))
        assert left is not None
        assert left.get(qn("w:color")) == "1A4D8F"

    def test_consecutive_border(self):
        """US1: 连续段各自有 pBdr."""
        doc = _convert(_HTML_QUICKSTART)
        for text in ("连续段一", "连续段二"):
            para = _find_para(doc, text)
            assert para is not None
            assert _pBdr(para) is not None, f"连续段 '{text}' pBdr 丢失"

    def test_paragraph_shading(self):
        """US1: background-color:#eef4fb → shd@fill=EEF4FB."""
        doc = _convert(_HTML_QUICKSTART)
        para = _find_para(doc, "浅蓝底纹段")
        assert para is not None
        shd = _shd(para)
        assert shd is not None, "shd 丢失"
        assert shd.get(qn("w:fill")) == "EEF4FB"

    def test_abstract_card_cell_shd_and_richtext(self):
        """US1-4: 裸 table 卡片 — cell shd + 多段落富文本完整."""
        doc = _convert(_HTML_QUICKSTART)
        # 找到 abstract-card table（第一个 table）
        assert len(doc.tables) >= 1
        card_cell = doc.tables[0].cell(0, 0)
        # cell shd
        tcPr = card_cell._tc.find(qn("w:tcPr"))
        assert tcPr is not None
        cell_shd = tcPr.find(qn("w:shd"))
        assert cell_shd is not None, "card cell shd 丢失"
        assert cell_shd.get(qn("w:fill")) == "EEF4FB"
        # 多段落
        texts = [p.text for p in card_cell.paragraphs if p.text.strip()]
        assert any("核心观点" in t for t in texts)
        assert any("第一段" in t for t in texts)
        assert any("第二段" in t for t in texts)
        # bold run
        assert any(
            r.font.bold for p in card_cell.paragraphs for r in p.runs
            if "核心观点" in p.text
        )

    def test_callout_richtext(self):
        """US3: callout 含 <strong> + 多段 → 单元格多段落 + bold run."""
        doc = _convert(_HTML_QUICKSTART)
        # callout 渲染为最后一个 table（1×1）
        callout_cell = doc.tables[-1].cell(0, 0)
        nonempty = [p for p in callout_cell.paragraphs if p.text.strip()]
        assert len(nonempty) >= 2, f"callout 段落不足: {len(nonempty)}"
        # 首段含 bold run
        assert any(r.font.bold for r in nonempty[0].runs), "callout 标题无 bold"
