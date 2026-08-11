"""Tests for html_to_docx.header_footer."""
from __future__ import annotations

import pytest
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from html_to_docx.header_footer import render_footer, render_header


def _tag(html: str, selector: str):
    return BeautifulSoup(html, "lxml").find(selector)


_HEADER_HTML = """
<header class="doc-header">
  <h1 class="doc-title">My Document Title</h1>
  <p class="doc-subtitle">Subtitle text here</p>
  <p class="doc-meta">Author: Test · Date: 2026-05-19</p>
</header>
"""

_FOOTER_HTML = """
<footer class="doc-footer">
  © 2026 Acme Corp · Confidential
</footer>
"""


def test_header_title():
    doc = Document()
    render_header(_tag(_HEADER_HTML, "header"), doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("My Document Title" in t for t in texts)


def test_header_subtitle():
    doc = Document()
    render_header(_tag(_HEADER_HTML, "header"), doc)
    for para in doc.paragraphs:
        if "Subtitle text here" in para.text:
            assert any(r.font.bold for r in para.runs if r.text.strip())
            break
    else:
        pytest.fail("Subtitle paragraph not found")


def test_header_meta():
    doc = Document()
    render_header(_tag(_HEADER_HTML, "header"), doc)
    for para in doc.paragraphs:
        if "Author: Test" in para.text:
            run = para.runs[0]
            assert run.font.size == Pt(9)
            break
    else:
        pytest.fail("Meta paragraph not found")


def test_footer_section():
    doc = Document()
    render_footer(_tag(_FOOTER_HTML, "footer"), doc)
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "Acme Corp" in footer_text


def test_footer_border_top():
    doc = Document()
    render_footer(_tag(_FOOTER_HTML, "footer"), doc)
    para = doc.sections[0].footer.paragraphs[0]
    assert "w:pBdr" in para._p.xml
    assert "w:top" in para._p.xml
