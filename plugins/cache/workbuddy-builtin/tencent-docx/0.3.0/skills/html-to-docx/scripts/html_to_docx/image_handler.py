"""image_handler.py — Embed images from three sources into a DOCX document.

Sources handled:
  1. data:image/...;base64,...  → decode and embed inline
  2. Local file path            → open and embed directly
  3. http/https URL             → download via httpx (5s timeout), or fallback text

Auto-scaling: if image width > max_width_cm, scale proportionally using Pillow.
Fallback: add paragraph "[图片无法加载]" or "[图片: <url>]" on error.
"""
from __future__ import annotations

import base64
import io
import os
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

_MAX_WIDTH_CM_DEFAULT = 16.0


def embed_image(img_tag, document: "Document", max_width_cm: float = _MAX_WIDTH_CM_DEFAULT) -> None:
    """Embed the image described by *img_tag* (a BeautifulSoup Tag) into *document*."""
    src: str = img_tag.get("src", "")
    width_attr = img_tag.get("width")
    height_attr = img_tag.get("height")

    if src.startswith("data:"):
        _embed_base64(src, document, max_width_cm, width_attr, height_attr)
    elif src.startswith("http://") or src.startswith("https://"):
        _embed_remote(src, document, max_width_cm, width_attr, height_attr)
    else:
        _embed_local(src, document, max_width_cm, width_attr, height_attr)


# ---------------------------------------------------------------------------
# Source handlers
# ---------------------------------------------------------------------------

def _embed_base64(
    src: str,
    document: "Document",
    max_width_cm: float,
    width_attr,
    height_attr,
) -> None:
    try:
        header, encoded = src.split(",", 1)
        image_bytes = base64.b64decode(encoded)
        _add_picture(document, io.BytesIO(image_bytes), max_width_cm, width_attr, height_attr)
    except Exception:  # noqa: BLE001
        document.add_paragraph("[图片无法加载]")


def _embed_local(
    src: str,
    document: "Document",
    max_width_cm: float,
    width_attr,
    height_attr,
) -> None:
    if not os.path.exists(src):
        document.add_paragraph("[图片无法加载]")
        return
    try:
        _add_picture(document, src, max_width_cm, width_attr, height_attr)
    except Exception:  # noqa: BLE001
        document.add_paragraph("[图片无法加载]")


def _embed_remote(
    src: str,
    document: "Document",
    max_width_cm: float,
    width_attr,
    height_attr,
) -> None:
    try:
        import httpx
        response = httpx.get(src, timeout=5.0, follow_redirects=True)
        response.raise_for_status()
        _add_picture(
            document, io.BytesIO(response.content), max_width_cm, width_attr, height_attr
        )
    except Exception:  # noqa: BLE001
        document.add_paragraph(f"[图片: {src}]")


# ---------------------------------------------------------------------------
# Picture insertion with auto-scaling
# ---------------------------------------------------------------------------

def _add_picture(document, image_source, max_width_cm: float, width_attr, height_attr) -> None:
    """Add picture to document, scaling if necessary."""
    from docx.shared import Cm

    # Determine display width
    width_cm: float | None = None
    if width_attr:
        try:
            width_px = float(str(width_attr).replace("px", "").strip())
            width_cm = width_px / 37.795  # 1cm ≈ 37.795px at 96dpi
        except ValueError:
            pass

    # Read image to check natural size via Pillow
    if isinstance(image_source, (str, os.PathLike)):
        with open(image_source, "rb") as f:
            image_bytes = f.read()
        image_source_for_docx: io.BytesIO | str = image_source  # type: ignore[assignment]
    else:
        # It's already a BytesIO; read and reset
        image_bytes = image_source.read()
        image_source.seek(0)
        image_source_for_docx = image_source

    natural_width_cm = _get_natural_width_cm(image_bytes)

    # Determine final width
    if width_cm is not None:
        final_width_cm = min(width_cm, max_width_cm)
    elif natural_width_cm is not None:
        final_width_cm = min(natural_width_cm, max_width_cm)
    else:
        final_width_cm = max_width_cm

    document.add_picture(image_source_for_docx, width=Cm(final_width_cm))


def _get_natural_width_cm(image_bytes: bytes) -> float | None:
    """Return natural image width in cm using Pillow, or None on failure."""
    try:
        from PIL import Image

        with Image.open(io.BytesIO(image_bytes)) as img:
            width_px, _ = img.size
            dpi = img.info.get("dpi", (96, 96))
            dpi_x = dpi[0] if isinstance(dpi, (tuple, list)) else dpi
            width_inch = width_px / dpi_x
            return width_inch * 2.54  # inches → cm
    except Exception:  # noqa: BLE001
        return None
