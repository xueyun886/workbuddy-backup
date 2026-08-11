"""section_model.py — Semantic <section> → docx section mapper.

Spec: S-26060126E (FR-001~007)

Maps top-level ``<section role="...">`` elements to docx sections, driving
implicit section breaks at section boundaries. Parses role / data-orientation /
data-margin-* (cm) / data-page-restart attributes and auto-assigns indices by
DOM order.

This module is the **sole writer** of ``<w:sectPr>`` (D-R5 single-writer):
``page_setup`` only produces ``PageDefaults``; ``apply_sections`` merges
defaults (base) with per-section overrides and injects sectPr.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Literal, Optional, Set


# Valid values for the orientation attribute.
_VALID_ORIENTATIONS = ("portrait", "landscape")

@dataclass
class SectionSpec:
    """Parsed result of a single top-level ``<section>``.

    Validation rules (data-model.md #1):
    - ``index``: assigned by converter in DOM order (0..N-1); MUST NOT come
      from template-authored ids (C-CON-006).
    - ``role``: open string, no hard-coded whitelist; used for named-page
      binding (``section[role=X]{page:Y}``).
    - ``orientation``: invalid value -> None + warning (falls back to portrait
      via PageDefaults).
    - ``margin_*`` (cm): parse failure (non-numeric) -> None (uses default),
      no error.
    - ``page_restart``: written to ``pgNumType@w:start``; non-positive integer
      -> None + warning.
    - ``page_name``: filled by ``page_css.bind_pages_to_sections`` from
      ``section[role=X]{page:Y}``; referenced ``@page <name>`` missing ->
      fallback default + warning.
    - ``explicit_attrs``: which attributes came from the ``<section>`` tag;
      drives per-property merge in ``apply_sections`` (D-R5 / q2).
    - ``boundary_marker``: invisible locator marker injected by
      ``mark_section_boundaries``; cleared after sectPr injection.
    """

    index: int
    role: Optional[str] = None
    orientation: Optional[Literal["portrait", "landscape"]] = None
    margin_top: Optional[float] = None       # cm
    margin_bottom: Optional[float] = None     # cm
    margin_left: Optional[float] = None       # cm
    margin_right: Optional[float] = None      # cm
    page_restart: Optional[int] = None
    page_name: Optional[str] = None           # filled by PageBinding
    explicit_attrs: Set[str] = field(default_factory=set)
    boundary_marker: Optional[str] = None     # internal locator, cleared post-apply


# Invisible boundary-marker prefix: a zero-width space + sentinel, carried
# through html4docx so apply_sections can locate each section's first paragraph.
_BOUNDARY_PREFIX = "\u200b__SEC_BOUNDARY__:"

# Block-level tags eligible to carry a boundary marker (text-bearing).
_BLOCK_TAGS = (
    "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "li", "blockquote", "pre", "div", "td", "th",
)


def _warn(msg: str) -> None:
    import warnings
    warnings.warn(msg, stacklevel=2)


def _parse_margin(value):  # type: ignore[no-untyped-def]
    """Parse a data-margin-* attribute (cm). Non-numeric -> None (no error)."""
    if value is None:
        return None
    try:
        return float(value)
    except (ValueError, TypeError):
        return None


def parse_sections(soup) -> "list[SectionSpec]":
    """Parse top-level ``<section>`` elements in DOM order (FR-001~006).

    - No ``<section>`` -> ``[SectionSpec(index=0)]`` (single, all None) for
      backward compatibility (FR-007 / C-CON-004).
    - Nested ``<section>`` -> only the top-level counts as a section; inner
      ones are demoted to plain blocks + warning (Edge Case).
    - Auto-assigns ``index`` (FR-006); never reads template-authored id/index
      (C-CON-006).
    """
    all_sections = soup.find_all("section")
    # Top-level only: a <section> whose ancestors include no other <section>.
    top_level = [s for s in all_sections if s.find_parent("section") is None]

    if not top_level:
        return [SectionSpec(index=0)]

    # Warn for nested sections (inner demoted to plain blocks).
    for s in all_sections:
        if s.find_parent("section") is not None:
            _warn(
                "Nested <section> detected; only the top-level <section> is "
                "treated as a section boundary, inner sections are demoted to "
                "plain blocks."
            )

    specs: list = []
    for idx, sec in enumerate(top_level):
        explicit: Set[str] = set()

        role = sec.get("role")
        role = role.strip() if isinstance(role, str) and role.strip() else None

        # orientation
        orientation = None
        raw_orient = sec.get("data-orientation")
        if raw_orient is not None:
            o = str(raw_orient).strip().lower()
            if o in _VALID_ORIENTATIONS:
                orientation = o
                explicit.add("orientation")
            else:
                _warn(
                    f"Invalid data-orientation {raw_orient!r}; ignoring "
                    f"(falls back to document default)."
                )

        # margins (per-property)
        margins = {}
        for side in ("top", "bottom", "left", "right"):
            raw = sec.get(f"data-margin-{side}")
            parsed = _parse_margin(raw)
            margins[side] = parsed
            if parsed is not None:
                explicit.add(f"margin_{side}")

        # page-restart
        page_restart = None
        raw_restart = sec.get("data-page-restart")
        if raw_restart is not None:
            try:
                n = int(str(raw_restart).strip())
                if n >= 1:
                    page_restart = n
                    explicit.add("page_restart")
                else:
                    _warn(
                        f"Invalid data-page-restart {raw_restart!r} "
                        f"(must be a positive integer); ignoring."
                    )
            except (ValueError, TypeError):
                _warn(
                    f"Invalid data-page-restart {raw_restart!r} "
                    f"(must be a positive integer); ignoring."
                )

        specs.append(
            SectionSpec(
                index=idx,
                role=role,
                orientation=orientation,  # type: ignore[arg-type]
                margin_top=margins["top"],
                margin_bottom=margins["bottom"],
                margin_left=margins["left"],
                margin_right=margins["right"],
                page_restart=page_restart,
                explicit_attrs=explicit,
            )
        )
    return specs


def mark_section_boundaries(soup, specs: "list[SectionSpec]") -> str:
    """Inject an invisible locator marker into each section's first block child.

    Returns the processed HTML so html4docx carries markers through; the actual
    sectPr injection is done later by ``apply_sections``. No-op (returns
    ``str(soup)``) for the single-section backward-compat case (no <section>).
    """
    all_sections = soup.find_all("section")
    top_level = [s for s in all_sections if s.find_parent("section") is None]
    if not top_level:
        # No <section> markup: nothing to mark (single-section path).
        return str(soup)

    import uuid
    for spec, sec in zip(specs, top_level):
        marker = f"{_BOUNDARY_PREFIX}{uuid.uuid4().hex[:12]}"
        spec.boundary_marker = marker
        first_block = _find_first_block(sec)
        if first_block is not None:
            # Prepend the marker text to the first block's text content so it
            # rides on the first paragraph that html4docx emits for this section.
            first_block.insert(0, marker)
        else:
            # Empty section: insert a paragraph carrying just the marker.
            p = soup.new_tag("p")
            p.string = marker
            sec.insert(0, p)
    return str(soup)


def _find_first_block(section):  # type: ignore[no-untyped-def]
    """Find the first element child of *section* if it is a block-level tag.

    Returns the element only when the **first** element child (skipping bare
    NavigableString whitespace) is in ``_BLOCK_TAGS``. If the section starts
    with a non-block element (e.g. ``<table>``), returns None so the caller
    inserts a dedicated marker ``<p>`` at the very beginning — ensuring the
    boundary marker always precedes all section content in
    ``document.paragraphs``.
    """
    from bs4 import NavigableString
    for el in section.children:
        # Skip whitespace-only text nodes
        if isinstance(el, NavigableString):
            if not el.strip():
                continue
            return None  # Non-empty text node as first content — unusual
        name = getattr(el, "name", None)
        if name in _BLOCK_TAGS:
            return el
        # First element child is not a block tag (e.g. <table>) — bail out
        return None
    return None


def apply_sections(document, specs: "list[SectionSpec]", defaults) -> None:
    """Inject sectPr at section boundaries; sole writer of ``<w:sectPr>`` (D-R5).

    Merge: ``defaults`` as base, ``SectionSpec`` as override.
    - Page size: always from ``defaults.width_cm/height_cm``; landscape swaps w/h.
    - orientation / margins: per-property — attr in ``explicit_attrs`` uses the
      section value, else falls back to ``defaults``.
    - ``page_restart`` -> ``sectPr/pgNumType@w:start`` (仅对 index 0 的单节生效)。
    - First section: no break. Sections 2..N: 方案 B 下降级为普通分页符
      （``pageBreakBefore``），不再插入分节符。
    - Boundary markers are removed after injection; trailing empty section
      skips the break.
    """
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # --- single-section backward-compat path (no markers) -----------------
    if len(specs) <= 1:
        spec = specs[0] if specs else SectionSpec(index=0)
        _apply_section_props(document.sections[0], spec, defaults)
        _set_page_restart(document.sections[0], spec, qn, OxmlElement)
        _clear_marker_in_doc(document, spec)
        return

    # --- multi-section path -----------------------------------------------
    # python-docx starts with one section (the body sectPr). We insert N-1
    # section breaks so the document ends up with N sections. The break for
    # section k is anchored at the first paragraph carrying spec[k]'s marker.
    body = document.element.body

    # Locate each section's first paragraph by its boundary marker.
    # NOTE: 用 casefold 匹配以兼容 CSS text-transform: uppercase 把哨兵大写化的场景
    # （详见 _clear_marker_in_doc 的坑点 2）。
    marker_para = {}
    for para in document.paragraphs:
        txt = para.text
        if not txt:
            continue
        txt_cf = txt.casefold()
        for spec in specs:
            if spec.boundary_marker and spec.boundary_marker.casefold() in txt_cf:
                marker_para[spec.index] = para
                break

    # Insert breaks for sections 1..N-1 (before each section's first paragraph).
    # Section 0 owns no break; its sectPr is the document's final sectPr.
    for spec in specs[1:]:
        anchor = marker_para.get(spec.index)
        if anchor is None:
            # Marker lost (e.g. empty section mangled): skip the break for this
            # section — degrade gracefully without aborting.
            continue
        _insert_section_break_before(anchor, body, qn, OxmlElement)

    # Re-resolve sections now that breaks exist; map spec.index -> section.
    sections = document.sections
    # The number of sections should equal number of breaks inserted + 1.
    # Apply per-section properties in order.
    for spec in specs:
        if spec.index < len(sections):
            _apply_section_props(sections[spec.index], spec, defaults)
            _set_page_restart(sections[spec.index], spec, qn, OxmlElement)

    # Remove all boundary markers from paragraph text.
    for spec in specs:
        _clear_marker_in_doc(document, spec)


def _insert_section_break_before(anchor_para, body, qn, OxmlElement) -> None:
    """方案 B：把节边界降级为普通「分页符」，不再产生分节符。

    原实现会在边界处插入一个挂着 ``<w:sectPr w:type="nextPage">`` 的空段落，
    Word 里会显示成一条丑陋的「分节符(下一页)」标记。这里改为直接在下一节的
    首段上设置 ``<w:pageBreakBefore/>``，让下一节照样另起一页，但**不产生分节符
    标记、也不产生多余空段落**。

    ⚠️ 代价：全文因此只有一个节，跨节的差异化页面设置（横向节 / 每节不同页边距 /
    页码从 1 重起 / 封面不显示页码）将**不再生效**——这些能力本质上依赖分节符，
    分页符无法实现。参数 ``body`` / ``qn`` / ``OxmlElement`` 保留仅为兼容调用签名。
    """
    anchor_para.paragraph_format.page_break_before = True


def _apply_section_props(section, spec: "SectionSpec", defaults) -> None:
    """Apply page size / orientation / margins to *section* (per-property merge)."""
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT

    # orientation: explicit section value, else defaults.
    if "orientation" in spec.explicit_attrs and spec.orientation:
        orientation = spec.orientation
    else:
        orientation = defaults.orientation

    # page size: always from defaults; swap for landscape.
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(defaults.height_cm)
        section.page_height = Cm(defaults.width_cm)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(defaults.width_cm)
        section.page_height = Cm(defaults.height_cm)

    # margins: per-property merge.
    section.top_margin = Cm(
        spec.margin_top if "margin_top" in spec.explicit_attrs and spec.margin_top is not None
        else defaults.margin_top
    )
    section.bottom_margin = Cm(
        spec.margin_bottom if "margin_bottom" in spec.explicit_attrs and spec.margin_bottom is not None
        else defaults.margin_bottom
    )
    section.left_margin = Cm(
        spec.margin_left if "margin_left" in spec.explicit_attrs and spec.margin_left is not None
        else defaults.margin_left
    )
    section.right_margin = Cm(
        spec.margin_right if "margin_right" in spec.explicit_attrs and spec.margin_right is not None
        else defaults.margin_right
    )


def _set_page_restart(section, spec: "SectionSpec", qn, OxmlElement) -> None:
    """Write ``sectPr/pgNumType@w:start`` for page-number restart (FR-005)."""
    if spec.page_restart is None:
        return
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(spec.page_restart))


def _clear_marker_in_doc(document, spec: "SectionSpec") -> None:
    """Strip a section's boundary marker from all paragraph runs.

    坑点 1（跨 run）：html4docx 常常把首段拆成多个 run（不同字体 / letter-spacing /
    零宽空格边界都会触发拆分），因此哨兵串会**横跨多个 run**。单 run 内
    ``marker in run.text`` 的精确匹配会失败。→ 这里以段落为单位把所有 run 文本拼起来
    定位哨兵，再按字符逐 run 抠掉。

    坑点 2（大小写）：CSS ``text-transform: uppercase``（常见于封面的 eyebrow 段）
    会把整段文本大写化，包括我们注入的哨兵里的 UUID hex 部分。此时
    ``spec.boundary_marker`` 存的还是小写，直接 substring 匹配会 False。→ 匹配统一
    用 ``str.casefold()`` 做大小写不敏感比对。哨兵串只含 ASCII，casefold 和 lower
    等价，不会引入 locale 相关的意外。
    """
    from docx.oxml.ns import qn as _qn

    marker = spec.boundary_marker
    if not marker:
        return
    marker_cf = marker.casefold()
    for para in document.paragraphs:
        if marker_cf not in para.text.casefold():
            continue
        _strip_substring_across_runs(para, marker)
        # If paragraph is now empty (no visible text) and has no sectPr,
        # remove it to avoid a spurious blank line.
        if not para.text.strip():
            p_el = para._p
            pPr = p_el.find(_qn("w:pPr"))
            has_sectpr = pPr is not None and pPr.find(_qn("w:sectPr")) is not None
            if not has_sectpr:
                p_el.getparent().remove(p_el)


def _strip_substring_across_runs(para, needle: str) -> None:
    """从段落里删除 *needle*，允许 needle 横跨多个连续 run，且大小写不敏感。

    做法：把所有 run 文本 casefold 后拼成 concat，在 concat 里定位 needle 的
    ``[start, end)``（同样 casefold），然后遍历各 run，用「字符区间相交」的方式，
    把落在区间内的字符从对应 run 里抠掉。

    ⚠️ 前提：``needle.casefold()`` 与 ``needle`` 长度一致（ASCII 时恒成立）。本模块
    哨兵是 ``\\u200b__SEC_BOUNDARY__:<12-hex>``，全部 ASCII，满足前提；对一般 needle
    则不适用。
    """
    runs = list(para.runs)
    if not runs:
        return
    concat_cf = "".join(r.text for r in runs).casefold()
    idx = concat_cf.find(needle.casefold())
    if idx < 0:
        return
    start, end = idx, idx + len(needle)

    cursor = 0
    for run in runs:
        rt = run.text
        if not rt:
            continue
        r_start, r_end = cursor, cursor + len(rt)
        cursor = r_end
        # 与哨兵区间 [start, end) 求交
        if r_end <= start or r_start >= end:
            continue  # 本 run 完全在哨兵外
        keep_left = rt[: max(0, start - r_start)]
        keep_right = rt[max(0, end - r_start):] if end < r_end else ""
        run.text = keep_left + keep_right
