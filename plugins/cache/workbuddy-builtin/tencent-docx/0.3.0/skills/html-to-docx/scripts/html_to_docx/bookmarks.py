"""HTML fillable-field parsing and DOCX bookmark materialization."""
from __future__ import annotations

import copy
import hashlib
import re
import uuid
from dataclasses import dataclass
from typing import Iterable

from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .types import FieldBinding

_FIELD_KEY_RE = re.compile(r"^[A-Za-z][A-Za-z0-9_-]{0,63}$")
_BOOKMARK_NAME_RE = re.compile(r"^[A-Za-z\u4e00-\u9fff][A-Za-z0-9_\u4e00-\u9fff]{0,39}$")
_TOKEN_PREFIX = "__DOCXFIELD_"


@dataclass(frozen=True)
class PendingField:
    key: str
    display_text: str
    token: str
    bookmark_name: str | None = None
    source: str = "explicit"
    label: str | None = None


def prepare_fields(html: str) -> tuple[str, list[PendingField], list[str]]:
    """Replace valid field spans with opaque text tokens before base conversion."""
    soup = BeautifulSoup(html, "lxml")
    fields: list[PendingField] = []
    warnings: list[str] = []
    seen: set[str] = set()
    seen_bookmark_names: set[str] = set()

    elements = soup.find_all(attrs={"data-docx-field": True})
    for element in elements:
        key = str(element.get("data-docx-field") or "")
        if not _FIELD_KEY_RE.fullmatch(key):
            warnings.append(f"Invalid docx field key: {key!r}")
            continue
        if key in seen:
            warnings.append(f"Duplicate docx field key: {key!r}")
            continue
        seen.add(key)
        requested_name = str(element.get("data-docx-bookmark") or "")
        bookmark_name = requested_name or None
        if bookmark_name and not _BOOKMARK_NAME_RE.fullmatch(bookmark_name):
            warnings.append(f"Invalid docx bookmark name: {bookmark_name!r}")
            bookmark_name = None
        if bookmark_name and bookmark_name in seen_bookmark_names:
            warnings.append(f"Duplicate docx bookmark name: {bookmark_name!r}")
            bookmark_name = None
        if bookmark_name:
            seen_bookmark_names.add(bookmark_name)
        token = f"{_TOKEN_PREFIX}{uuid.uuid4().hex}__"
        display_text = element.get_text() or "________________"
        element.clear()
        element.append(token)
        fields.append(
            PendingField(
                key=key,
                display_text=display_text,
                token=token,
                bookmark_name=bookmark_name,
            )
        )

    _append_automatic_fields(soup, fields, warnings)
    return str(soup), fields, warnings


def materialize_fields(document, fields: Iterable[PendingField]) -> tuple[list[FieldBinding], list[str]]:
    """Replace field tokens in DOCX paragraphs with paired Word bookmarks."""
    bindings: list[FieldBinding] = []
    warnings: list[str] = []
    next_id = _next_bookmark_id(document)
    used_names = {
        element.get(qn("w:name"))
        for element in document.element.iter(qn("w:bookmarkStart"))
        if element.get(qn("w:name"))
    }

    for field in fields:
        matching = [para for para in _iter_paragraphs(document) if field.token in para.text]
        if len(matching) != 1:
            warnings.append(f"Docx field could not be safely located: {field.key!r}")
            continue
        name = field.bookmark_name or _bookmark_name(field.key, used_names)
        if name in used_names:
            warnings.append(f"Duplicate docx bookmark name in document: {name!r}")
            continue
        if _replace_token_with_bookmark(matching[0], field.token, field.display_text, name, next_id):
            bindings.append(
                FieldBinding(field.key, name, field.display_text, field.source, field.label)
            )
            used_names.add(name)
            next_id += 1
        else:
            warnings.append(f"Docx field could not be safely materialized: {field.key!r}")

    return bindings, warnings


def _iter_paragraphs(document) -> Iterable[Paragraph]:
    seen: set = set()

    def visit_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para._p not in seen:
                        seen.add(para._p)
                        yield para
                for nested in cell.tables:
                    yield from visit_table(nested)

    for para in document.paragraphs:
        if para._p not in seen:
            seen.add(para._p)
            yield para
    for table in document.tables:
        yield from visit_table(table)


def _next_bookmark_id(document) -> int:
    ids = [
        int(value)
        for element in document.element.iter(qn("w:bookmarkStart"))
        if (value := element.get(qn("w:id"))) and value.isdigit()
    ]
    return max(ids, default=-1) + 1


def _bookmark_name(key: str, used_names: set[str]) -> str:
    normalized = key.replace("-", "_")
    candidate = f"fld_{normalized}"
    if len(candidate) > 40 or candidate in used_names:
        candidate = f"fld_{hashlib.sha256(key.encode()).hexdigest()[:32]}"
    return candidate


def _append_automatic_fields(soup, fields: list[PendingField], warnings: list[str]) -> None:
    """Recognize common Chinese contract blanks in text and table cells."""
    used_labels: dict[str, int] = {}

    def add_field(element, label: str, replacement: str) -> None:
        if not label:
            warnings.append("Automatic docx field has no Chinese label")
            return
        used_labels[label] = used_labels.get(label, 0) + 1
        number = used_labels[label]
        token = f"{_TOKEN_PREFIX}{uuid.uuid4().hex}__"
        key = f"auto_{len(fields) + 1}"
        name = f"{label}_{number:02d}"
        element.string = replacement + token
        fields.append(PendingField(key, replacement or "\u00a0", token, name, "automatic", label))

    for cell in soup.find_all(["td", "th"]):
        if cell.find(attrs={"data-docx-field": True}):
            continue
        if not cell.get_text(strip=True):
            row = cell.find_parent("tr")
            cells = row.find_all(["td", "th"], recursive=False) if row else []
            index = cells.index(cell) if cell in cells else 0
            if index > 0:
                label = cells[index - 1].get_text(" ", strip=True).rstrip("：:")
                add_field(cell, label, "")

    pattern = re.compile(r"(?P<label>[\u4e00-\u9fff]{2,12})[：:](?P<line>[_＿—-]{3,})|(?P<label2>[\u4e00-\u9fff]{2,12})(?P<line2>[_＿—-]{3,})")
    for text_node in list(soup.find_all(string=True)):
        parent = text_node.parent
        if (
            parent is None
            or parent.name in {"style", "script"}
            or parent.find_parent(["td", "th"]) is not None
            or parent.find(attrs={"data-docx-field": True})
        ):
            continue
        text = str(text_node)
        match = pattern.search(text)
        if match is None:
            continue
        label = match.group("label") or match.group("label2")
        prefix = text[:match.start()]
        visible = text[match.start():match.end() - len(match.group("line") or match.group("line2") or "")]
        holder = soup.new_tag("span")
        holder.string = prefix
        text_node.replace_with(holder)
        add_field(holder, label, visible)


def _replace_token_with_bookmark(
    paragraph: Paragraph,
    token: str,
    display_text: str,
    bookmark_name: str,
    bookmark_id: int,
) -> bool:
    runs = list(paragraph.runs)
    joined = "".join(run.text for run in runs)
    start = joined.find(token)
    if start < 0 or joined.find(token, start + len(token)) >= 0:
        return False
    end = start + len(token)

    spans: list[tuple[int, int]] = []
    cursor = 0
    for run in runs:
        next_cursor = cursor + len(run.text)
        spans.append((cursor, next_cursor))
        cursor = next_cursor

    affected = [index for index, (left, right) in enumerate(spans) if left < end and right > start]
    if not affected:
        return False
    first_index, last_index = affected[0], affected[-1]
    first_left, _ = spans[first_index]
    last_left, _ = spans[last_index]
    before = runs[first_index].text[: start - first_left]
    after = runs[last_index].text[end - last_left :]
    parent = paragraph._p
    original_index = parent.index(runs[first_index]._r)
    before_run = copy.deepcopy(runs[first_index]._r)
    after_run = copy.deepcopy(runs[last_index]._r)

    for run in runs[first_index : last_index + 1]:
        parent.remove(run._r)

    elements = []
    if before:
        elements.append(_run_with_text(before_run, before))
    elements.append(_bookmark_start(bookmark_id, bookmark_name))
    elements.append(_run_with_text(before_run, display_text))
    elements.append(_bookmark_end(bookmark_id))
    if after:
        elements.append(_run_with_text(after_run, after))
    for offset, element in enumerate(elements):
        parent.insert(original_index + offset, element)
    return True


def _run_with_text(run, text: str):
    for node in list(run.findall(qn("w:t"))):
        run.remove(node)
    text_node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def _bookmark_start(bookmark_id: int, name: str):
    element = OxmlElement("w:bookmarkStart")
    element.set(qn("w:id"), str(bookmark_id))
    element.set(qn("w:name"), name)
    return element


def _bookmark_end(bookmark_id: int):
    element = OxmlElement("w:bookmarkEnd")
    element.set(qn("w:id"), str(bookmark_id))
    return element
