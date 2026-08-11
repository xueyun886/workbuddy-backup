"""converter.py — Core HTML→DOCX conversion orchestrator.

Pipeline:
  Phase 0  css_resolver.resolve()          — expand CSS var()
  Phase 0a image_resolver.resolve_image_paths() — resolve relative img src to absolute
  Phase 0b style_injector.inject_styles()  — CSS selector → inline style
  Phase 0c text_splitter.split_cjk_latin() — split CJK/Latin into spans
  Phase 1  extract special regions          — replace with placeholders to preserve order
  Phase 2  html-for-docx base conversion   — standard HTML elements (incl. placeholders)
  Phase 2b apply_tag_styles_to_document    — tag CSS → docx built-in styles
  Phase 2c apply_cjk_run_properties        — eastAsia font hints on CJK runs
  Phase 2d apply_table_styles              — table borders, cell shading, cell text styles
  Phase 3  replace placeholders             — swap placeholder paragraphs with real content
  Phase 4  page_setup                      — size / orientation / margins
  Phase 5  save                             — write .docx to output_path
"""
from __future__ import annotations

import os
import tempfile
import uuid
import warnings

from collections.abc import Callable

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from html4docx import HtmlToDocx

from .css_resolver import resolve as css_resolve
from .page_setup import resolve_page_defaults
from .types import ConvertOptions, ConvertResult

# Prefix used to identify placeholder paragraphs in the docx output
_PLACEHOLDER_PREFIX = "\u200b__SPECIAL_REGION__:"


class Converter:
    """Orchestrates the full HTML→DOCX pipeline."""

    def __init__(self, options: ConvertOptions | None = None) -> None:
        self.options = options or ConvertOptions()
        self._warnings: list[str] = []
        self._fields = []

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def convert(self, html: str) -> ConvertResult:
        """Convert *html* string to a .docx file.

        Returns a :class:`ConvertResult` with ``success=True`` and
        ``docx_path`` set on success, or ``success=False`` and
        ``error`` / ``markdown_fallback`` on failure.
        """
        try:
            docx_path = self._run_pipeline(html)
            return ConvertResult(
                success=True,
                docx_path=docx_path,
                warnings=list(self._warnings),
                fields=list(self._fields),
            )
        except Exception as exc:  # noqa: BLE001
            from .fallback import html_to_markdown
            return ConvertResult(
                success=False,
                error=str(exc),
                markdown_fallback=html_to_markdown(html),
                warnings=list(self._warnings),
            )

    # ------------------------------------------------------------------
    # Pipeline implementation
    # ------------------------------------------------------------------

    def _run_pipeline(self, html: str) -> str:
        """Execute the conversion pipeline; return the output file path."""

        # Phase 0 — CSS variable resolution
        html = css_resolve(html)

        # Phase 0a — Resolve relative image paths to absolute (based on base_dir)
        if self.options.base_dir:
            from .image_resolver import resolve_image_paths
            html = resolve_image_paths(html, self.options.base_dir)

        # Phase 0b — CSS selector parsing + class inlining + tag styles extraction
        from .style_injector import inject_styles, apply_tag_styles_to_document
        html, tag_styles = inject_styles(html)

        # Phase 0c — Split mixed CJK/Latin text into separate spans
        from .text_splitter import split_cjk_latin
        html = split_cjk_latin(html)

        # Capture the styled HTML *before* boundary markers / placeholders are
        # injected, so page_css / page_setup parse clean CSS + meta.
        clean_html = html

        # Phase 0c' — Replace fillable field spans with stable text tokens.
        # html4docx does not retain arbitrary data-* attributes, so the tokens
        # are materialized as bookmarks only after all document mutations finish.
        from .bookmarks import prepare_fields
        html, pending_fields, field_warnings = prepare_fields(html)
        for message in field_warnings:
            self._warn(message)

        # Phase 1 — Replace special regions with placeholders (preserves order)
        html, region_map = self._extract_special_regions(html)

        # Phase 2e (pre) — Section parsing + boundary marking (D-R2).
        # Parse <section> on the placeholder-substituted soup, then inject
        # invisible markers so html4docx carries them through to docx paragraphs.
        from .section_model import parse_sections, mark_section_boundaries
        section_soup = BeautifulSoup(html, "lxml")
        section_specs = parse_sections(section_soup)
        html = mark_section_boundaries(section_soup, section_specs)

        # Phase 1b — Strip insignificant whitespace between block elements.
        # html4docx treats whitespace text nodes (newlines + indentation)
        # between container/block tags as paragraph content, generating
        # spurious empty paragraphs. Remove them before feeding to html4docx.
        html = self._strip_inter_block_whitespace(html)

        # Phase 2 — Base conversion via html-for-docx
        document = Document()
        parser = HtmlToDocx()

        try:
            parser.add_html_to_document(html, document)
        except Exception as exc:  # noqa: BLE001
            self._warn(f"html-for-docx base conversion warning: {exc}")

        # Phase 2b — Apply tag selector styles to docx built-in styles
        apply_tag_styles_to_document(document, tag_styles)

        # Phase 2c — Apply CJK run properties for correct East Asian rendering
        from .style_injector import apply_cjk_run_properties
        apply_cjk_run_properties(document, tag_styles)

        # Phase 2d — Apply table styles (borders, cell shading, cell text styles)
        from .table_style_applier import apply_table_styles
        apply_table_styles(document, html)

        # Phase 2d (post) — Apply body paragraph styles (border/shd/em spacing).
        # 命门集成 (T020b / FR-001~003)：正文普通 <p>/<div>/<h*>/<li> 走 html4docx
        # 从不调用 apply_paragraph_styles → 段落边框/底纹/em 段距端到端丢失。
        # 镜像 apply_table_styles 的后处理范式，复用 US1/US2 的 apply_paragraph_styles。
        # 用 clean_html（内联后、placeholder 替换前快照）以拿到完整 inline style；
        # document.paragraphs 天然只含 body 顶层段落（不含表格单元格段落），与
        # table_style_applier 天然隔离、不重复处理表格内段落。
        from .body_paragraph_styler import apply_body_paragraph_styles
        apply_body_paragraph_styles(document, clean_html)

        # Phase 3 — Replace placeholder paragraphs with real rendered content
        self._replace_placeholders(document, region_map)

        # Phase 2e (post) — Section break + sectPr injection (sole writer, D-R5).
        # Replaces the old standalone apply_page_setup call.
        from .section_model import apply_sections
        defaults = resolve_page_defaults(clean_html, self.options)
        apply_sections(document, section_specs, defaults)

        # Phase 2e (post) — CSS @page furniture (header/footer + dynamic fields).
        # Parse the bounded @page subset from the clean (pre-marker) HTML, bind
        # named pages to section roles, then render margin-boxes into each
        # section's header/footer; finally arm field auto-refresh (T018/T026 /
        # FR-008~018 / D-R3).
        from .page_css import (
            parse_page_css,
            bind_pages_to_sections,
            apply_page_furniture,
        )
        from .fields import enable_update_fields

        page_model = parse_page_css(clean_html)
        bind_pages_to_sections(page_model, section_specs)
        apply_page_furniture(document, page_model, section_specs)
        enable_update_fields(document)

        # Phase 4 — Remove spurious empty paragraphs (html4docx artifact).
        # html4docx generates empty <w:p> for whitespace between container
        # tags and for <div> open tags. Remove them.
        self._strip_spurious_empty_paragraphs(document)

        # Phase 4b — Materialize fillable field tokens after every existing
        # operation that can alter paragraphs, runs, tables, or sections.
        from .bookmarks import materialize_fields
        self._fields, field_warnings = materialize_fields(document, pending_fields)
        for message in field_warnings:
            self._warn(message)

        # Phase 5 — Save
        output_path = self._resolve_output_path()
        document.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Pre-processing: strip insignificant whitespace between blocks
    # ------------------------------------------------------------------

    @staticmethod
    def _strip_inter_block_whitespace(html: str) -> str:
        """Remove whitespace-only text nodes between block-level elements.

        html4docx (HTMLParser-based) treats any text between tags as paragraph
        content. Newlines and indentation between block elements like
        ``</div>\\n    <h2>`` or ``<div>\\n    <p>`` produce empty paragraphs.

        This method uses BeautifulSoup to find and extract NavigableString
        nodes that are pure whitespace and sit between two block-level
        siblings or as the first/last child of a block container.
        """
        from bs4 import BeautifulSoup, NavigableString

        _BLOCK_TAGS = frozenset((
            "div", "p", "h1", "h2", "h3", "h4", "h5", "h6",
            "ul", "ol", "li", "table", "tr", "td", "th", "thead", "tbody",
            "section", "blockquote", "pre", "nav", "header", "footer",
            "figure", "figcaption", "main", "article", "aside",
        ))

        soup = BeautifulSoup(html, "lxml")
        # Collect whitespace-only NavigableStrings that are between blocks
        to_remove = []
        for text_node in soup.find_all(string=True):
            if not isinstance(text_node, NavigableString):
                continue
            if text_node.strip():
                continue  # Has visible content — keep

            parent = text_node.parent
            if parent is None:
                continue
            parent_name = getattr(parent, "name", None)
            if parent_name not in _BLOCK_TAGS:
                continue  # Inside inline element — keep

            # Check siblings: is this between two block elements or at
            # the edge of a block container?
            prev_sib = text_node.previous_sibling
            next_sib = text_node.next_sibling

            prev_is_block = (
                prev_sib is None
                or getattr(prev_sib, "name", None) in _BLOCK_TAGS
            )
            next_is_block = (
                next_sib is None
                or getattr(next_sib, "name", None) in _BLOCK_TAGS
            )

            if prev_is_block and next_is_block:
                to_remove.append(text_node)

        for node in to_remove:
            node.extract()

        return str(soup)

    # ------------------------------------------------------------------
    # Post-processing: strip spurious empty paragraphs
    # ------------------------------------------------------------------

    @staticmethod
    def _strip_spurious_empty_paragraphs(document) -> None:
        """Remove spurious empty paragraphs generated by html4docx.

        html4docx creates empty <w:p> elements for:
        1. Whitespace text nodes between block tags
        2. <div> open tags (creates a new paragraph on entering any div)

        This method removes empty paragraphs (no visible w:t text, no
        sectPr) that appear:
        - At the start of the document body (leading empties)
        - Between two content-bearing elements (paragraphs with text or
          tables) — single empty paragraphs that serve no formatting purpose

        Preserves: paragraphs with sectPr (section breaks), consecutive
        empty paragraphs that might be intentional spacing (keeps max 1
        removal between content elements).
        """
        from docx.oxml.ns import qn

        body = document.element.body
        children = list(body)

        def _is_empty_p(el):
            """True if element is <w:p> with no visible text and no sectPr."""
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag != "p":
                return False
            has_text = any(
                (t.text or "").strip()
                for t in el.iter(qn("w:t"))
            )
            if has_text:
                return False
            # A paragraph carrying an image / drawing / embedded object is NOT
            # empty even without <w:t> text. html4docx emits <img> as a run
            # containing only a <w:drawing> (no text); treating it as an empty
            # paragraph and removing it would orphan the media part and drop the
            # picture from the rendered document.
            for media_tag in ("w:drawing", "w:pict", "w:object"):
                if next(el.iter(qn(media_tag)), None) is not None:
                    return False
            pPr = el.find(qn("w:pPr"))
            has_sectpr = pPr is not None and pPr.find(qn("w:sectPr")) is not None
            return not has_sectpr

        def _is_content_p(el):
            """True if element is a non-empty paragraph (not table)."""
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag == "p":
                return not _is_empty_p(el)
            return False

        to_remove = []
        i = 0
        n = len(children)

        # Pass 1: Remove leading empty paragraphs
        while i < n and _is_empty_p(children[i]):
            to_remove.append(children[i])
            i += 1

        # Pass 2: Remove empty paragraphs only between two non-empty paragraphs.
        # Preserve empty paragraphs adjacent to tables (they serve as spacing).
        while i < n:
            if _is_empty_p(children[i]):
                # Look ahead: find next non-empty element
                j = i + 1
                while j < n and _is_empty_p(children[j]):
                    j += 1
                # Only remove if BOTH neighbors are non-empty paragraphs
                prev_is_p = i > 0 and _is_content_p(children[i - 1])
                next_is_p = j < n and _is_content_p(children[j])
                if prev_is_p and next_is_p:
                    for k in range(i, j):
                        to_remove.append(children[k])
                i = j if j > i else i + 1
            else:
                i += 1

        for el in to_remove:
            body.remove(el)

    # ------------------------------------------------------------------
    # Special region extraction — replace with placeholders
    # ------------------------------------------------------------------

    def _extract_special_regions(self, html: str) -> tuple[str, dict]:
        """Replace special regions with placeholder <p> tags.

        Returns (modified_html, region_map) where region_map maps
        placeholder_id → {"type": ..., "html": ...}.
        """
        soup = BeautifulSoup(html, "lxml")
        region_map: dict[str, dict] = {}
        header_html: str | None = None
        footer_html: str | None = None

        # TODO: OOXML 和 html 在Header 标签上存在语义冲突，后续解决，目前按照正文走html4docx渲染（@heulwenmai)
        # Extract <header> (no class) — "文档封面区"，渲染为正文最前面的普通段落
        # header = soup.find("header")
        # if header and not header.get("class"):
        #     header_html = str(header)
        #     header.decompose()

        # TODO: 支持 class="doc-header" 写入 OOXML section.header（页眉，每页顶部重复）
        # doc_header = soup.find(class_="doc-header")
        # if doc_header:
        #     region_map["__doc_header__"] = {"type": "ooxml-header", "html": str(doc_header)}
        #     doc_header.decompose()

        # Extract <footer class="doc-footer"> — goes to OOXML section.footer（页脚，每页底部重复）
        footer = soup.find("footer", class_="doc-footer")
        if footer:
            footer_html = str(footer)
            footer.decompose()

        # Replace special components (data-component="...") in document order.
        comps = soup.find_all(attrs={"data-component": True})

        for el in comps:
            pid = self._make_placeholder_id()
            region_map[pid] = {"type": "component", "html": str(el)}
            el.replace_with(self._make_placeholder_tag(soup, pid))

        # Store header/footer in region_map under special keys (non-body)
        if header_html:
            region_map["__header__"] = {"type": "header", "html": header_html}
        if footer_html:
            region_map["__footer__"] = {"type": "footer", "html": footer_html}

        return str(soup), region_map

    def _replace_placeholders(
        self, document: DocxDocument, region_map: dict[str, dict]
    ) -> None:
        """Find placeholder paragraphs and replace with rendered content."""
        from .header_footer import render_header, render_footer
        from .toc_renderer import render_toc
        from .components import render_component

        # Handle header/footer (not positional — goes to docx sections)
        header_info = region_map.pop("__header__", None)
        footer_info = region_map.pop("__footer__", None)

        if footer_info:
            ftr = BeautifulSoup(footer_info["html"], "lxml").find("footer")
            if ftr:
                render_footer(ftr, document)

        # Handle body-level placeholders (toc / components) — order preserved
        if not region_map and not header_info:
            return

        # Snapshot paragraphs BEFORE render_header so we don't capture header
        # elements that haven't been moved yet (render_header appends then
        # repositions, so the paragraph list must be taken after it settles).
        if header_info:
            hdr = BeautifulSoup(header_info["html"], "lxml").find("header")
            if hdr:
                render_header(hdr, document)

        if not region_map:
            return

        # Walk paragraphs to find placeholders and replace in-place.
        # Snapshot AFTER header insertion so the paragraph list is stable.
        resolved_pids: set[str] = set()
        paragraphs = list(document.paragraphs)
        # print("Before render placeholder: paragraphs", paragraphs)
        for para in paragraphs:
            text = para.text
            if not text.startswith(_PLACEHOLDER_PREFIX):
                continue

            pid = text[len(_PLACEHOLDER_PREFIX):]
            info = region_map.get(pid)
            if info is None:
                continue

            if info["type"] == "toc":
                nav = BeautifulSoup(info["html"], "lxml").find("nav")
                if nav:
                    self._render_at_paragraph(
                        document, para, lambda doc, _nav=nav: render_toc(_nav, doc)
                    )
                    resolved_pids.add(pid)
            elif info["type"] == "component":
                comp_soup = BeautifulSoup(info["html"], "lxml")
                el = comp_soup.find(attrs={"data-component": True})
                if el is not None:
                    self._render_at_paragraph(
                        document, para, lambda doc, _el=el, _para=para: render_component(_el, doc, _para)
                    )
                    resolved_pids.add(pid)

        # Warn about any unresolved placeholders (html4docx may have mangled them)
        for pid, info in region_map.items():
            if pid not in resolved_pids:
                self._warn(
                    f"Special region placeholder not found in docx output: "
                    f"type={info['type']}, id={pid}"
                )

    def _render_at_paragraph(self, document: DocxDocument, placeholder_para: Paragraph, render_fn: Callable[[DocxDocument], None]) -> None:
        """Replace a placeholder paragraph with content generated by render_fn.

        Strategy: render_fn appends content to the document end, then we move
        those new paragraphs/tables to the placeholder's position and remove
        the placeholder.

        Important: python-docx's add_table() inserts a trailing empty <w:p>
        after every table, and also manipulates the <w:sectPr> node at the very
        end of body.  We therefore re-query the body children *after* render_fn
        completes and use the actual parent check before removal.
        """
        body = document.element.body
        # NOTE: _element is python-docx's de-facto public API for accessing
        # the underlying lxml element of a paragraph.
        placeholder_elem = placeholder_para._element

        # Render — 在占位段落前插入组件，需要组件渲染方法实现插入
        render_fn(document)

        if placeholder_elem.getparent() is body:
            body.remove(placeholder_elem)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _make_placeholder_id() -> str:
        return uuid.uuid4().hex[:12]

    @staticmethod
    def _make_placeholder_tag(soup: BeautifulSoup, pid: str) -> Tag:
        """Create a <p> tag that will become a recognizable paragraph in docx."""
        tag = soup.new_tag("p")
        tag.string = f"{_PLACEHOLDER_PREFIX}{pid}"
        return tag

    def _resolve_output_path(self) -> str:
        if self.options.output_path:
            return self.options.output_path
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        return path

    def _warn(self, msg: str) -> None:
        self._warnings.append(msg)
        warnings.warn(msg, stacklevel=3)
