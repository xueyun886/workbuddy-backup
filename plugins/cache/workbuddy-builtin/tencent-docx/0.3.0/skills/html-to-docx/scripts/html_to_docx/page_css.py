"""page_css.py — CSS @page bounded-subset parser + page furniture renderer.

Spec: S-26060126E (FR-008~013)

Parses the bounded ``@page`` subset (D-R1): ``@page`` / ``@page <name>``,
6 margin-boxes, content token combos, ``string-set``, and
``section[role=X]{page:Y}`` bindings. Renders margin-boxes into per-section
header/footer paragraphs with Tab-stop layout (D-SP-12), named-page binding,
and OOXML headerReference inheritance (D-R4).

Unsupported subset syntax → ignored + warning (never aborts; FR-013 / SC-009).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from enum import Enum
from typing import Dict, List, Optional

from .fields import PageContentToken


class MarginBoxPos(str, Enum):
    """The 6 supported CSS Paged Media margin-box positions (FR-009)."""

    TOP_LEFT = "top-left"
    TOP_CENTER = "top-center"
    TOP_RIGHT = "top-right"
    BOTTOM_LEFT = "bottom-left"
    BOTTOM_CENTER = "bottom-center"
    BOTTOM_RIGHT = "bottom-right"


@dataclass
class PageRule:
    """One ``@page`` / ``@page <name>`` rule.

    - ``name``: None = default ``@page``; non-None = named page ``@page <name>``.
    - ``margin_boxes``: pos -> ordered content tokens; non-supported positions
      are ignored + warning at parse time.
    """

    name: Optional[str] = None
    margin_boxes: Dict[MarginBoxPos, List[PageContentToken]] = field(default_factory=dict)


@dataclass
class StringSetBinding:
    """``<selector> { string-set: <name> content(text) }`` binding.

    Stores the **raw selector string** (e.g. ``'h1'``) only; the
    selector -> styleId resolution happens at field-write time in
    ``fields.resolve_styleref_style_id`` (FR-017). ``page_css`` carries no
    style-mapping knowledge.
    """

    name: str
    selector: str


@dataclass
class PageCssModel:
    """Parsed result of the ``@page`` bounded subset for a document.

    - ``rules``: page name (None=default) -> PageRule.
    - ``page_bindings``: role -> page_name (from ``section[role=X]{page:Y}``).
    - ``string_sets``: list of StringSetBinding (raw selectors).
    """

    rules: Dict[Optional[str], PageRule] = field(default_factory=dict)
    page_bindings: Dict[str, str] = field(default_factory=dict)
    string_sets: List[StringSetBinding] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Parsing (D-R1: bounded regex; no CSS library)
# ---------------------------------------------------------------------------

import re

from .fields import tokenize_content, TokenKind

# Map margin-box at-keyword (without leading @) to MarginBoxPos.
_MARGIN_BOX_NAMES = {pos.value: pos for pos in MarginBoxPos}

# @page <name>? { ... }  — capture optional name + body (with nested { }).
# We cannot match nested braces with plain regex, so we scan manually.
_AT_PAGE_RE = re.compile(r"@page\b([^\{]*)\{", re.IGNORECASE)

# @<pos> { <decls> } inside a @page body.
_MARGIN_BOX_RE = re.compile(
    r"@([a-z-]+)\s*\{([^}]*)\}", re.IGNORECASE | re.DOTALL
)

# content: <value> ;
_CONTENT_DECL_RE = re.compile(
    r"content\s*:\s*(.*?)\s*(?:;|$)", re.IGNORECASE | re.DOTALL
)

# selector { string-set: <name> content(text) }
_STRING_SET_RE = re.compile(
    r"([^{}@]+?)\s*\{[^}]*?string-set\s*:\s*([A-Za-z_][\w-]*)\s+content\(\s*text\s*\)[^}]*\}",
    re.IGNORECASE | re.DOTALL,
)

# section[role=X] { page: Y }
_PAGE_BINDING_RE = re.compile(
    r"section\s*\[\s*role\s*=\s*[\"']?([\w-]+)[\"']?\s*\]\s*\{[^}]*?page\s*:\s*([\w-]+)[^}]*\}",
    re.IGNORECASE | re.DOTALL,
)

# Supported @page selectors: bare "@page" and "@page <name>".  Anything with
# a pseudo-class (e.g. ":first") or extra tokens is unsupported.
_PAGE_NAME_RE = re.compile(r"^\s*([A-Za-z_][\w-]*)?\s*$")


def _warn(msg: str) -> None:
    import warnings
    warnings.warn(msg, stacklevel=2)


def _extract_style_text(html: str) -> str:
    """Concatenate all <style> block contents; strip comments."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    parts = []
    for style_tag in soup.find_all("style"):
        css = style_tag.string or style_tag.get_text() or ""
        css = re.sub(r"/\*.*?\*/", "", css, flags=re.DOTALL)
        parts.append(css)
    return "\n".join(parts)


def _match_balanced_block(text: str, open_pos: int) -> int:
    """Given index of an opening '{', return index just after its matching '}'.

    Returns -1 if unbalanced.
    """
    depth = 0
    i = open_pos
    n = len(text)
    while i < n:
        c = text[i]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return i + 1
        i += 1
    return -1


def _parse_margin_boxes(body: str) -> "dict":
    """Parse the 6 supported margin-boxes from a @page body."""
    boxes: dict = {}
    for m in _MARGIN_BOX_RE.finditer(body):
        raw_pos = m.group(1).strip().lower()
        decls = m.group(2)
        pos = _MARGIN_BOX_NAMES.get(raw_pos)
        if pos is None:
            _warn(f"Unsupported @page margin-box {raw_pos!r}; ignoring.")
            continue
        cm = _CONTENT_DECL_RE.search(decls)
        if cm is None:
            continue
        raw_content = cm.group(1).strip()
        # Empty string literal "" is not a legal value (FR-026).
        if raw_content in ('""', "''"):
            _warn(
                "margin-box content empty string \"\" is not a legal value; "
                "treated as unset (inherits previous section)."
            )
            continue
        boxes[pos] = tokenize_content(raw_content)
    return boxes


def parse_page_css(html: str) -> PageCssModel:
    """Parse the bounded ``@page`` subset from a document's <style> blocks (D-R1).

    Supports: ``@page`` / ``@page <name>``, 6 margin-boxes, the 5 content token
    kinds, ``string-set``, and ``section[role=X]{page:Y}``. Unsupported subset
    syntax is ignored + warning (FR-013); never aborts (SC-009).
    """
    model = PageCssModel()
    css = _extract_style_text(html)

    # --- @page rules (manual scan for balanced braces) --------------------
    pos = 0
    for m in _AT_PAGE_RE.finditer(css):
        header = m.group(1).strip()
        brace_open = m.end() - 1  # index of '{'
        end = _match_balanced_block(css, brace_open)
        if end == -1:
            _warn("Unbalanced @page block; ignoring.")
            continue
        body = css[brace_open + 1:end - 1]

        # Validate the @page selector against the supported subset.
        name_match = _PAGE_NAME_RE.match(header)
        if name_match is None:
            _warn(
                f"Unsupported @page selector {header!r} (e.g. :first pseudo-class "
                f"or size declarations); ignoring this @page rule."
            )
            continue
        name = name_match.group(1)  # None for default @page
        boxes = _parse_margin_boxes(body)
        rule = PageRule(name=name, margin_boxes=boxes)
        model.rules[name] = rule

    # --- string-set bindings ----------------------------------------------
    for m in _STRING_SET_RE.finditer(css):
        selector = m.group(1).strip()
        name = m.group(2).strip()
        # Skip @page bodies accidentally matched (selector starting with @).
        if selector.startswith("@") or "{" in selector:
            continue
        model.string_sets.append(StringSetBinding(name=name, selector=selector))

    # --- section[role=X]{page:Y} bindings ---------------------------------
    for m in _PAGE_BINDING_RE.finditer(css):
        role = m.group(1).strip()
        page_name = m.group(2).strip()
        model.page_bindings[role] = page_name

    return model


def bind_pages_to_sections(model: PageCssModel, specs: "list") -> None:
    """Fill each ``SectionSpec.page_name`` from ``page_bindings`` (FR-011).

    - role with no binding -> page_name stays None (uses default @page).
    - bound page_name with no matching PageRule -> fallback default + warning.
    """
    for spec in specs:
        role = spec.role
        if not role:
            continue
        page_name = model.page_bindings.get(role)
        if page_name is None:
            continue
        if page_name not in model.rules:
            _warn(
                f"section[role={role!r}] binds to @page {page_name!r} which is "
                f"not defined; falling back to default @page."
            )
            spec.page_name = None
            continue
        spec.page_name = page_name


# ---------------------------------------------------------------------------
# Rendering page furniture into docx headers/footers (D-SP-12 / D-R4 / FR-009/012)
# ---------------------------------------------------------------------------

# Which margin-boxes belong to the header (top) vs footer (bottom) region, and
# their horizontal alignment slot (left / center / right) within that region.
_TOP_BOXES = (
    (MarginBoxPos.TOP_LEFT, "left"),
    (MarginBoxPos.TOP_CENTER, "center"),
    (MarginBoxPos.TOP_RIGHT, "right"),
)
_BOTTOM_BOXES = (
    (MarginBoxPos.BOTTOM_LEFT, "left"),
    (MarginBoxPos.BOTTOM_CENTER, "center"),
    (MarginBoxPos.BOTTOM_RIGHT, "right"),
)


def _rule_for_spec(model: PageCssModel, spec) -> Optional[PageRule]:
    """Resolve the effective ``PageRule`` for a section (named page or default)."""
    page_name = getattr(spec, "page_name", None)
    if page_name is not None and page_name in model.rules:
        return model.rules[page_name]
    return model.rules.get(None)


def _region_tokens(rule: Optional[PageRule], boxes) -> "dict":
    """Collect declared margin-box token lists for one region (top/bottom).

    Returns ``{slot: tokens}`` for slots that are **declared** in the rule.
    Missing margin-boxes are simply absent (the missing/none/empty tri-state:
    missing -> not in dict; ``none`` -> present but all-NONE tokens; empty
    string was already dropped at parse time -> behaves as missing).
    """
    region: dict = {}
    if rule is None:
        return region
    for pos, slot in boxes:
        if pos in rule.margin_boxes:
            region[slot] = rule.margin_boxes[pos]
    return region


def _region_is_empty(region: "dict") -> bool:
    """True when a region has no declared box, or every declared box is all-NONE.

    All-NONE means "no furniture": the part must not be created (封面真正无家具,
    FR-026). A missing region is likewise empty.
    """
    if not region:
        return True
    for tokens in region.values():
        for tok in tokens:
            if tok.kind != TokenKind.NONE:
                return False
    return True


def _content_emkey(rule: Optional[PageRule]) -> str:
    """A stable identity key for a PageRule's furniture content.

    Two sections sharing the same furniture content reuse one header/footer
    part (linked inheritance). Page name alone is insufficient because the
    default @page is shared by many sections; we key on the rendered token
    structure so consecutive sections collapse onto one part (SC-002).
    """
    if rule is None:
        return "<none>"
    parts = []
    for pos in MarginBoxPos:
        toks = rule.margin_boxes.get(pos)
        if not toks:
            continue
        sig = ",".join(f"{t.kind.value}:{t.value or ''}" for t in toks)
        parts.append(f"{pos.value}={sig}")
    return "|".join(parts) if parts else "<none>"


def _twips(cm: float) -> int:
    """Convert centimetres to twips (1 cm = 567 twips)."""
    return int(round(cm * 567))


def _section_widths_cm(section):  # type: ignore[no-untyped-def]
    """Return (content_width_cm, left_margin_cm) for tab-stop math.

    Content width = page_width - left_margin - right_margin. Falls back to
    sensible A4 defaults when python-docx returns None.
    """
    from docx.shared import Cm

    def _cm(emu, default):
        if emu is None:
            return default
        return emu / Cm(1)

    page_w = _cm(section.page_width, 21.0)
    left = _cm(section.left_margin, 3.17)
    right = _cm(section.right_margin, 3.17)
    content_w = page_w - left - right
    if content_w <= 0:
        content_w = page_w
    return content_w, left


def _write_region_paragraph(paragraph, region, content_width_cm, styleref_resolver):  # type: ignore[no-untyped-def]
    """Render a top/bottom region into a single paragraph with Tab stops (D-SP-12).

    Layout: left slot at the line start (left-aligned), center slot at a centre
    tab stop (content-width midpoint), right slot at a right tab stop (right
    margin). Slots are separated by ``\\t``. The absolute tab positions derive
    from the section's content width (FR-009).
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm
    from .fields import write_field

    # Configure tab stops: centre at midpoint, right at content-width end.
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(content_width_cm / 2.0), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(content_width_cm), WD_TAB_ALIGNMENT.RIGHT)

    # Emit slots in left, center, right order, separated by tabs. A tab is
    # written before center/right slots so they land on their tab stops even
    # when an earlier slot is absent.
    wrote_any = False
    for slot in ("left", "center", "right"):
        if slot in ("center", "right"):
            paragraph.add_run("\t")
        tokens = region.get(slot)
        if tokens:
            write_field(paragraph, tokens, styleref_resolver)
            wrote_any = True
    return wrote_any


def apply_page_furniture(document, model: PageCssModel, specs: "list") -> None:
    """Render margin-boxes into each section's header/footer (D-SP-12 / D-R4).

    - Same-region boxes (top/bottom) merge into a **single** header/footer
      paragraph with Tab stops: left at line start, center at the content-width
      midpoint, right at the right margin; ``\\t`` between slots (D-SP-12 /
      FR-009).
    - A margin-box that is **missing** -> not written (the part inherits the
      previous section via OOXML; R5).
    - Consecutive sections sharing the same PageRule furniture -> the first
      writes the part, later sections stay ``linked_to_previous`` and inherit
      the same rId (D-R4 / SC-002). Inheritance is content-keyed, so it holds
      across orientation/page-width differences without recomputing tab stops
      (FR-012 / D-SP-12 補充).
    - ``content: none`` -> that box produces nothing; a region whose boxes are
      all-NONE -> no part is created/linked (封面真正无家具, FR-026).
    - ``content: ""`` (illegal) -> dropped at parse time, behaves as missing
      (warning already emitted; inherits previous section, FR-026).
    - Field tokens are delegated to ``fields.write_field``.
    """
    from .fields import resolve_styleref_style_id

    sections = document.sections
    string_sets = model.string_sets

    def _resolver(name):  # bound resolver -> styleId | None
        return resolve_styleref_style_id(name, string_sets)

    # Track the last content key per region so consecutive sections sharing the
    # same furniture inherit (linked) rather than re-writing the part.
    last_header_key: Optional[str] = None
    last_footer_key: Optional[str] = None

    for spec in specs:
        idx = getattr(spec, "index", None)
        if idx is None or idx >= len(sections):
            continue
        section = sections[idx]
        rule = _rule_for_spec(model, spec)
        content_width_cm, _left = _section_widths_cm(section)

        top_region = _region_tokens(rule, _TOP_BOXES)
        bottom_region = _region_tokens(rule, _BOTTOM_BOXES)
        key = _content_emkey(rule)

        # --- header (top) -------------------------------------------------
        if _region_is_empty(top_region):
            # No furniture: leave linked so an empty default header is shown.
            last_header_key = "<none>"
        elif key == last_header_key:
            # Same furniture as the previous section -> inherit (link) the part.
            section.header.is_linked_to_previous = True
        else:
            header = section.header
            header.is_linked_to_previous = False
            para = header.paragraphs[0]
            _clear_paragraph(para)
            _write_region_paragraph(para, top_region, content_width_cm, _resolver)
            last_header_key = key

        # --- footer (bottom) ----------------------------------------------
        if _region_is_empty(bottom_region):
            last_footer_key = "<none>"
        elif key == last_footer_key:
            section.footer.is_linked_to_previous = True
        else:
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0]
            _clear_paragraph(para)
            _write_region_paragraph(para, bottom_region, content_width_cm, _resolver)
            last_footer_key = key


def _clear_paragraph(paragraph) -> None:  # type: ignore[no-untyped-def]
    """Remove all existing runs/field children from a paragraph in place."""
    p_el = paragraph._p
    for child in list(p_el):
        # Keep paragraph properties (w:pPr); strip content children.
        tag = child.tag
        if tag.endswith("}pPr"):
            continue
        p_el.remove(child)
