"""Tests for html_to_docx.toc_renderer."""
from __future__ import annotations

import pytest
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm

from html_to_docx.toc_renderer import render_toc


def _nav(html: str):
    return BeautifulSoup(html, "lxml").find("nav")


_SIMPLE_TOC = """
<nav class="doc-toc">
  <ol>
    <li><a href="#h1">Heading One</a>
      <ol>
        <li><a href="#h2">Heading Two</a>
          <ol>
            <li><a href="#h3">Heading Three</a></li>
          </ol>
        </li>
      </ol>
    </li>
  </ol>
</nav>
"""

_PAGED_TOC = """
<nav class="doc-toc">
  <ol>
    <li data-page="1"><a href="#intro">Introduction</a></li>
    <li data-page="5"><a href="#body">Body</a></li>
  </ol>
</nav>
"""


def test_three_level_toc():
    doc = Document()
    render_toc(_nav(_SIMPLE_TOC), doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("Heading One" in t for t in texts)
    assert any("Heading Two" in t for t in texts)
    assert any("Heading Three" in t for t in texts)


def test_indent_correct():
    doc = Document()
    render_toc(_nav(_SIMPLE_TOC), doc)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    # Level 1 has no indent; level 2 has Cm(0.5)
    level1 = next(p for p in paras if "Heading One" in p.text)
    level2 = next(p for p in paras if "Heading Two" in p.text)
    l1_indent = level1.paragraph_format.left_indent or 0
    l2_indent = level2.paragraph_format.left_indent or 0
    assert l2_indent > l1_indent


def test_link_to_text():
    doc = Document()
    render_toc(_nav(_SIMPLE_TOC), doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Should not contain href="..." artifacts
    assert not any("href" in t for t in texts)
    assert any("Heading One" in t for t in texts)


def test_page_number():
    doc = Document()
    render_toc(_nav(_PAGED_TOC), doc)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("1" in t and "Introduction" in t for t in texts)
    assert any("5" in t and "Body" in t for t in texts)
