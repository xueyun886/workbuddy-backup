"""Data card component renderer.

Renders ``data-component="data-card"`` as a 1×1 table with colored top border.

Actual HTML structure (produced by doc-typeset):
  <div data-component="data-card" data-title="标题" data-color="success">
    <span class="card-value">≤ 55</span>
    <span class="card-unit">亿元</span>
    <span class="card-trend card-trend--up">↑ 12.5%</span>
    <span class="card-label">描述文字</span>
  </div>

Also supports KV mode (data-type="kv"):
  <div data-component="data-card" data-title="项目状态" data-type="kv">
    <dl class="card-kv-list">
      <dt>进度</dt><dd>75%</dd>
    </dl>
  </div>
"""
from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.table import _Cell
from bs4 import Tag

from . import register
from ..style_mapper import apply_run_styles, apply_paragraph_styles
from ._cell_richtext import render_inline_into_paragraph

# Color variants — maps data-color to (border hex, value text hex)
_COLOR_MAP: dict[str, tuple[str, str]] = {
    "primary": ("333333", "111827"),
    "success": ("2E7D32", "2E7D32"),
    "warning": ("E65100", "E65100"),
    "danger":  ("C62828", "C62828"),
    "info":    ("01579B", "01579B"),
}


def _hex_to_rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _apply_top_border(cell: "_Cell", color_hex: str) -> None:
    """Apply a colored top border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "18")  # ~2.25pt
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color_hex)
    tcBorders.append(top)
    tcPr.append(tcBorders)


def render_stat_card_in_cell(element: Tag, cell: "_Cell") -> None:
    """Render a stat-card into a given table cell (used by data_card_grid too).

    Iterates direct children **in their original HTML order**. Each child is
    dispatched by its class name to a specific renderer; unknown elements get
    a generic paragraph render.
    """
    title = element.get("data-title", "")
    color_name = element.get("data-color", "primary").lower()
    border_hex, value_hex = _COLOR_MAP.get(color_name, _COLOR_MAP["primary"])

    # Top border
    _apply_top_border(cell, border_hex)

    # Use first default paragraph for title
    first_para = cell.paragraphs[0]
    first_para.text = ""

    if title:
        run = first_para.add_run(title)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Iterate children in document order
    for child in element.children:
        # --- Bare text node ---
        if isinstance(child, str):
            text = child.strip()
            if text:
                para = cell.add_paragraph()
                run = para.add_run(text)
                # 继承父元素的 inline style
                parent_style = element.get("style", "")
                if parent_style:
                    apply_run_styles(run, parent_style)
                    apply_paragraph_styles(para, parent_style)
            continue

        if not isinstance(child, Tag):
            continue

        child_classes = set(child.get("class") or [])

        # --- card-value (+ immediately look for sibling card-unit to append) ---
        if "card-value" in child_classes:
            para = cell.add_paragraph()
            run = para.add_run(child.get_text(strip=True))
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = _hex_to_rgb(value_hex)
            # Unit is visually part of value — peek at next sibling
            next_sib = child.find_next_sibling()
            if next_sib and isinstance(next_sib, Tag) and "card-unit" in set(next_sib.get("class") or []):
                run = para.add_run(f" {next_sib.get_text(strip=True)}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                # Mark unit as consumed so it's skipped in the next iteration
                next_sib["data-_rendered"] = "1"
            continue

        # --- card-unit (skip if already consumed by card-value above) ---
        if "card-unit" in child_classes:
            if child.get("data-_rendered"):
                continue
            # Standalone unit (no preceding card-value)
            para = cell.add_paragraph()
            run = para.add_run(child.get_text(strip=True))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        # --- card-trend ---
        if "card-trend" in child_classes:
            para = cell.add_paragraph()
            run = para.add_run(child.get_text(strip=True))
            run.font.size = Pt(9)
            run.font.bold = True
            if "card-trend--up" in child_classes:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            elif "card-trend--down" in child_classes:
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
            continue

        # --- card-label ---
        if "card-label" in child_classes:
            para = cell.add_paragraph()
            run = para.add_run(child.get_text(strip=True))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        # --- card-kv-list (dl with dt/dd pairs) ---
        if "card-kv-list" in child_classes:
            dts = child.find_all("dt")
            dds = child.find_all("dd")
            for dt, dd in zip(dts, dds):
                para = cell.add_paragraph()
                run_k = para.add_run(f"{dt.get_text(strip=True)}: ")
                run_k.font.size = Pt(9)
                run_k.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run_v = para.add_run(dd.get_text(strip=True))
                run_v.font.bold = True
            continue

        # --- Generic element: unknown class or no class ---
        text = child.get_text(strip=True)
        if text:
            para = cell.add_paragraph()
            # 行内样式注入 — 复用共享 cell 富文本渲染范式（D-R2，禁双轨）
            style_attr = child.get("style", "")
            render_inline_into_paragraph(child, para, style_attr)
            if not para.runs:
                # 兜底：helper 未产出 run（极端空白）时维持原行为
                para.add_run(text)
            if style_attr:
                apply_paragraph_styles(para, style_attr)


@register("data-card")
def render_data_card(element: Tag, document: DocxDocument, anchor: Paragraph | None = None) -> None:
    """Render a data-card component into *document*."""
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    render_stat_card_in_cell(element, cell)

    if anchor is not None:
        # 如果需要对表格进行分离，需要在表格前插入一个空段落，不然会被office识别成一整个表格
        # anchor._element.add_p_before()
        anchor._element.addprevious(table._element)
