"""_cell_richtext.py — Shared cell rich-text rendering helper.

Renders a bs4 ``Tag``'s children into a docx ``_Cell`` as independent
paragraphs / runs (preserving ``<strong>``, multiple ``<p>``, ``<li>`` instead
of flattening to text). Shared by ``data_card`` and ``callout`` (D-R2).

Exports:
  - ``render_children_into_cell(element, cell) -> None``
  - ``render_inline_into_paragraph(node, paragraph, parent_style="")`` (helper)

Extracted from the generic "iterate children → cell paragraphs/runs + inherit
inline style" logic of ``data_card.py`` (D-R2 / FR-005). Card-specific class
branches (``card-value`` / ``card-trend`` / …) are intentionally **not** part
of this shared helper — they remain in ``data_card``.
"""
from __future__ import annotations

from docx.table import _Cell
from docx.text.paragraph import Paragraph
from bs4 import NavigableString, Tag

from ..style_mapper import apply_run_styles, apply_paragraph_styles

# Inline tags that map to run-level formatting rather than new paragraphs.
_BOLD_TAGS = {"strong", "b"}
_ITALIC_TAGS = {"em", "i"}
_UNDERLINE_TAGS = {"u", "ins"}
_INLINE_TAGS = _BOLD_TAGS | _ITALIC_TAGS | _UNDERLINE_TAGS | {
    "span", "a", "code", "small", "sub", "sup", "mark",
}
# Block tags that each become an independent paragraph.
_BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li"}


def _is_empty_paragraph(paragraph: "Paragraph") -> bool:
    """A freshly added cell paragraph has no runs and empty text."""
    return not paragraph.runs and not paragraph.text.strip()


def render_inline_into_paragraph(
    node, paragraph: "Paragraph", parent_style: str = "", inherit_bold: bool = False,
    inherit_italic: bool = False, inherit_underline: bool = False,
) -> None:
    """Render an inline *node* (text or inline Tag) into *paragraph* as run(s).

    Recursively walks inline children, mapping ``<strong>``/``<b>`` → bold,
    ``<em>``/``<i>`` → italic, ``<u>`` → underline, and inheriting the
    surrounding inline formatting flags.
    """
    if isinstance(node, NavigableString):
        text = str(node)
        if not text.strip():
            return
        run = paragraph.add_run(text)
        if inherit_bold:
            run.font.bold = True
        if inherit_italic:
            run.font.italic = True
        if inherit_underline:
            run.font.underline = True
        if parent_style:
            apply_run_styles(run, parent_style)
        return

    if not isinstance(node, Tag):
        return

    name = (node.name or "").lower()
    bold = inherit_bold or name in _BOLD_TAGS
    italic = inherit_italic or name in _ITALIC_TAGS
    underline = inherit_underline or name in _UNDERLINE_TAGS
    own_style = node.get("style", "") or parent_style

    if not list(node.children):
        # Self-closing / empty inline tag (e.g. <br>) — emit nothing meaningful.
        return

    for child in node.children:
        render_inline_into_paragraph(
            child, paragraph, own_style, bold, italic, underline,
        )


def _add_block_paragraph(element: "Tag", cell: "_Cell", style_attr: str = "") -> "Paragraph":
    """Append a paragraph to *cell* (reusing the empty default first paragraph)."""
    first = cell.paragraphs[0]
    if _is_empty_paragraph(first) and len(cell.paragraphs) == 1:
        para = first
    else:
        para = cell.add_paragraph()
    if style_attr:
        apply_paragraph_styles(para, style_attr)
    return para


def render_children_into_cell(element: "Tag", cell: "_Cell") -> None:
    """Render *element*'s children into *cell* as independent paragraphs/runs.

    - Bare (non-empty) text node → its own paragraph + run.
    - Inline tag (``<strong>``/``<em>``/``<span>`` …) at top level → coalesced
      into the current/last paragraph as run(s) with formatting.
    - Block tag (``<p>``/``<div>``/``<h*>``/``<blockquote>``) → independent
      paragraph; inline children become runs (``<strong>`` → bold run).
    - ``<ul>``/``<ol>`` → each ``<li>`` becomes its own paragraph.
    - Nested ``<table>`` → degraded to a paragraph with its text (never crash).
    - Inline style on each node is injected via apply_run/paragraph_styles.
    """
    parent_style = element.get("style", "") or ""

    for child in element.children:
        # --- Bare text node ---
        if isinstance(child, NavigableString):
            text = str(child)
            if not text.strip():
                continue
            para = _add_block_paragraph(element, cell, parent_style)
            run = para.add_run(text)
            if parent_style:
                apply_run_styles(run, parent_style)
            continue

        if not isinstance(child, Tag):
            continue

        name = (child.name or "").lower()

        # --- Lists: each <li> is an independent paragraph ---
        if name in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                para = _add_block_paragraph(element, cell, child.get("style", ""))
                for sub in li.children:
                    render_inline_into_paragraph(sub, para, li.get("style", ""))
                if _is_empty_paragraph(para):
                    para.add_run(li.get_text(strip=True))
            continue

        # --- Nested table: degrade to text paragraph, never crash ---
        if name == "table":
            text = child.get_text(separator=" ", strip=True)
            if text:
                para = _add_block_paragraph(element, cell, child.get("style", ""))
                para.add_run(text)
            continue

        # --- Inline-level tag at top level → coalesce into current paragraph ---
        if name in _INLINE_TAGS:
            # Reuse the last paragraph if it already has content, else a new one.
            if len(cell.paragraphs) >= 1 and not _is_empty_paragraph(cell.paragraphs[-1]):
                para = cell.paragraphs[-1]
            else:
                para = _add_block_paragraph(element, cell, parent_style)
            render_inline_into_paragraph(child, para, child.get("style", "") or parent_style)
            continue

        # --- Block-level tag → independent paragraph with inline runs ---
        style_attr = child.get("style", "")
        para = _add_block_paragraph(element, cell, style_attr)
        has_content = False
        for sub in child.children:
            before = len(para.runs)
            render_inline_into_paragraph(sub, para, style_attr)
            if len(para.runs) > before:
                has_content = True
        if not has_content:
            text = child.get_text(strip=True)
            if text:
                para.add_run(text)
                if style_attr:
                    apply_run_styles(para.runs[-1], style_attr)
