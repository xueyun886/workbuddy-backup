"""table_style_applier.py — Phase 2d: Apply table styles from HTML to docx tables.

After html4docx base conversion (Phase 2), tables exist in the docx but lack:
  - border-color / border-style from <table> or <td> inline styles
  - background-color (cell shading) from <td>/<th> inline styles
  - font-size / color from <td>/<th> inline styles applied to cell paragraphs

This module re-parses the original HTML to extract table/cell styles, then
applies them to the corresponding docx tables via python-docx XML manipulation.
"""
from __future__ import annotations

import re
from typing import TYPE_CHECKING

from bs4 import BeautifulSoup, Tag
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from ._ooxml_borders import (
    BORDER_STYLE_MAP as _BORDER_STYLE_MAP,
    build_border_element,
    parse_border_shorthand as _parse_border_shorthand,
    parse_color as _parse_color,
)

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import Table, _Cell


# ---------------------------------------------------------------------------
# CSS parsing helpers (reused from style_mapper)
# ---------------------------------------------------------------------------

def _parse_style(style_attr: str) -> dict[str, str]:
    """Parse inline style string into property→value dict."""
    props: dict[str, str] = {}
    for decl in style_attr.split(";"):
        decl = decl.strip()
        if ":" in decl:
            k, _, v = decl.partition(":")
            props[k.strip().lower()] = v.strip()
    return props


def _parse_size_to_pt(value: str) -> float | None:
    """Convert CSS size string to points."""
    value = value.strip().lower()
    m = re.match(r"^([\d.]+)(px|pt|rem|em|cm|mm)?$", value)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2) or "pt"
    conversions = {"pt": num, "px": num * 0.75, "rem": num * 12, "em": num * 12,
                   "cm": num * 28.3465, "mm": num * 2.83465}
    return conversions.get(unit, num)


# ---------------------------------------------------------------------------
# Border helpers
# ---------------------------------------------------------------------------

_BORDER_SIDES = ("top", "bottom", "left", "right")


def _extract_first_row_col_pcts(html_table) -> list[float | None]:
    """Return per-column CSS width percentages from the HTML table's first row.

    Reads ``<td>/<th style="width: X%">`` (also honors an explicit ``width``
    attribute like ``width="25%"``). Returns a list of floats (in 0..100) or
    ``None`` for columns without a percentage width.
    """
    first_row = html_table.find("tr")
    if first_row is None:
        return []
    cells = first_row.find_all(["td", "th"], recursive=False)
    pcts: list[float | None] = []
    for cell in cells:
        pct: float | None = None
        style = cell.get("style", "")
        if style:
            props = _parse_style(style)
            wv = props.get("width", "").strip()
            if wv.endswith("%"):
                try:
                    pct = float(wv[:-1])
                except ValueError:
                    pct = None
        if pct is None:
            wattr = (cell.get("width") or "").strip()
            if wattr.endswith("%"):
                try:
                    pct = float(wattr[:-1])
                except ValueError:
                    pct = None
        pcts.append(pct)
    return pcts


def _normalize_cell_widths_to_pct(table: Table, html_table) -> None:
    """Rewrite per-cell ``<w:tcW>`` (and ``<w:tblGrid>``) as percentages.

    html4docx maps ``<td style="width: 25%">`` to ``<w:tcW w:type="dxa"
    w:w="500"/>`` — it drops the percent sign and multiplies by 20, producing
    absurdly small fixed widths (500 twips ≈ 0.88 cm). Combined with our
    ``tblLayout=autofit`` override, Word still respects the per-cell dxa
    values and squeezes columns to a single character wide, forcing
    character-per-line vertical stacking in each cell.

    This helper:
      1. Extracts per-column CSS width % from the HTML first row (falls back
         to equal split when a column has no percentage).
      2. Rewrites every ``<w:tc>``'s ``<w:tcW>`` to ``type="pct"`` with the
         corresponding value (``pct = X * 50`` where 5000 = 100%).
      3. Rewrites ``<w:tblGrid>/<w:gridCol>`` widths proportionally so Word's
         first-pass layout also reflects the intended distribution.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl

    # 1) Column percentages from HTML.
    pcts_raw = _extract_first_row_col_pcts(html_table)
    n_cols = len(tbl.findall(qn("w:tblGrid") + "/" + qn("w:gridCol")))
    if n_cols == 0:
        # Fall back to first tr's cell count if tblGrid is missing (rare).
        first_tr = tbl.find(qn("w:tr"))
        n_cols = len(first_tr.findall(qn("w:tc"))) if first_tr is not None else len(pcts_raw)
    if n_cols == 0:
        return

    # Align pcts length with actual column count; fill missing with None.
    pcts: list[float | None] = list(pcts_raw[:n_cols])
    while len(pcts) < n_cols:
        pcts.append(None)

    # If none of the columns provide a %, do NOT normalize — leave html4docx's
    # widths intact and let tblLayout=autofit distribute proportionally to
    # the (still small) tcW values, avoiding surprises on tables where the
    # author intentionally relied on content-based sizing.
    if not any(p is not None for p in pcts):
        # But we still must neutralize the pathological case where all tcW
        # values are absurdly small (< 720 twips ≈ 1.27cm), which means
        # html4docx got no width hints at all and just produced tiny defaults.
        # In that case, distribute equally.
        first_tr = tbl.find(qn("w:tr"))
        if first_tr is None:
            return
        tc_ws = [tc.find(qn("w:tcPr") + "/" + qn("w:tcW")) for tc in first_tr.findall(qn("w:tc"))]
        if all(
            (w is not None
             and w.get(qn("w:type")) == "dxa"
             and int(w.get(qn("w:w")) or "0") < 720)
            for w in tc_ws
        ):
            pcts = [100.0 / n_cols] * n_cols
        else:
            return

    # Fill Nones with equal share of the remainder.
    known_sum = sum(p for p in pcts if p is not None)
    n_unknown = sum(1 for p in pcts if p is None)
    if n_unknown:
        remaining = max(0.0, 100.0 - known_sum)
        share = remaining / n_unknown if n_unknown else 0.0
        pcts = [share if p is None else p for p in pcts]

    # Normalize so pcts sum to 100 (guard against >100 or <100 inputs).
    total = sum(pcts) or 100.0
    pcts = [p * 100.0 / total for p in pcts]

    # 2) Rewrite each <w:tc>'s <w:tcW>.
    for tr in tbl.findall(qn("w:tr")):
        tcs = tr.findall(qn("w:tc"))
        for col_idx, tc in enumerate(tcs):
            if col_idx >= n_cols:
                break
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:type"), "pct")
            tc_w.set(qn("w:w"), str(int(round(pcts[col_idx] * 50))))

    # 3) Rewrite <w:tblGrid>/<w:gridCol> proportionally, using the total
    #    of existing gridCol widths as the reference so we don't have to
    #    reach into section geometry here.
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        grid_cols = tbl_grid.findall(qn("w:gridCol"))
        # Use max(sum-of-current-widths, a sane default) as scaling base.
        current_sum = sum(int(gc.get(qn("w:w")) or "0") for gc in grid_cols)
        base = current_sum if current_sum >= 720 * len(grid_cols) else 8640  # ~15.24cm
        for col_idx, gc in enumerate(grid_cols):
            if col_idx >= n_cols:
                break
            gc.set(qn("w:w"), str(int(round(base * pcts[col_idx] / 100.0))))


def _force_table_full_width(table: Table, html_table) -> None:
    """Force table width to 100% of available page width.

    Sets ``<w:tblW w:type="pct" w:w="5000"/>`` (100% in fifths of a percent)
    and changes ``tblLayout`` to ``autofit`` so Word distributes column widths
    proportionally. This fixes html4docx's default behavior of producing
    ``tblW=0/auto`` with tiny fixed gridCol widths.

    Additionally normalizes per-cell widths from html4docx's ``tcW dxa=X*20``
    (a mistranslation of ``width: X%``) to ``tcW pct=X*50``, so Word honors
    the intended percentage distribution instead of collapsing columns to a
    single character wide. See ``_normalize_cell_widths_to_pct``.

    Skips tables that have an explicit narrow width CSS (e.g. ``width: 50%``
    or ``width: 200px``).
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Check if HTML explicitly sets a non-100% width
    style = html_table.get("style", "")
    if style:
        props = _parse_style(style)
        width_val = props.get("width", "").strip()
        if width_val and width_val != "100%":
            # Explicit narrow width — don't force full width
            return

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    # Set tblW to 100% (5000 fifths-of-a-percent)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    # Change layout to autofit so Word distributes column widths
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")

    # Normalize per-cell widths from html4docx's mistranslated dxa values.
    _normalize_cell_widths_to_pct(table, html_table)


def _strip_cell_leading_empty_paras(table: Table) -> None:
    """Remove leading empty paragraphs from all cells in *table*.

    html4docx generates an empty paragraph for whitespace/newlines between
    ``<td>`` and its first child element (e.g. ``<td>\\n    <p>``). This
    appears as a blank line before the actual cell content. We remove such
    paragraphs only when the cell has more than one paragraph (to avoid
    producing an invalid empty cell with zero paragraphs).
    """
    from docx.oxml.ns import qn

    for row in table.rows:
        for cell in row.cells:
            paras = cell.paragraphs
            if len(paras) <= 1:
                continue
            first = paras[0]
            # Check: no visible text in first paragraph
            has_text = any(
                (t.text or "").strip()
                for t in first._p.iter(qn("w:t"))
            )
            if not has_text:
                cell._tc.remove(first._p)


def _set_table_borders(table: Table, props: dict[str, str]) -> None:
    """Set border on all sides of a docx able from CSS properties."""
    # Collect border info from various CSS properties
    border_info: dict[str, dict[str, str]] = {}

    # border (shorthand for all sides)
    if "border" in props:
        parsed = _parse_border_shorthand(props["border"])
        for side in _BORDER_SIDES:
            border_info[side] = dict(parsed)

    # border-color (applies to all sides)
    if "border-color" in props:
        color = _parse_color(props["border-color"])
        if color:
            for side in _BORDER_SIDES:
                border_info.setdefault(side, {})["color"] = props["border-color"]

    # border-style
    if "border-style" in props:
        for side in _BORDER_SIDES:
            border_info.setdefault(side, {})["style"] = props["border-style"]

    # Per-side overrides
    for side in _BORDER_SIDES:
        side_key = f"border-{side}"
        if side_key in props:
            parsed = _parse_border_shorthand(props[side_key])
            border_info.setdefault(side, {}).update(parsed)
        if f"border-{side}-color" in props:
            border_info.setdefault(side, {})["color"] = props[f"border-{side}-color"]
        if f"border-{side}-style" in props:
            border_info.setdefault(side, {})["style"] = props[f"border-{side}-style"]

    if not border_info:
        return

    # Apply to OOXML
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    # Also add insideH and insideV for internal borders
    all_sides = list(_BORDER_SIDES) + ["insideH", "insideV"]

    for side in all_sides:
        info = border_info.get(side) or border_info.get("top", {})  # fallback to first defined
        if not info:
            continue

        # Build via shared border utility (preserves legacy table color 000000).
        elem = build_border_element(side, info, default_color="000000")

        # Remove existing element for this side if any
        existing = tblBorders.find(qn(f"w:{side}"))
        if existing is not None:
            tblBorders.remove(existing)
        tblBorders.append(elem)


def _set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """Set background color (shading) on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)

def _padding_shorthand_sides(value: str) -> dict[str, str]:
    """Expand a CSS ``padding`` shorthand into top/right/bottom/left raw values.

    Follows CSS 1–4 value rules (mirrors margin shorthand semantics).
    Returns a dict with keys top/right/bottom/left (raw strings).
    """
    parts = value.split()
    if not parts:
        return {}
    if len(parts) == 1:
        t = r = b = l = parts[0]
    elif len(parts) == 2:
        t = b = parts[0]
        r = l = parts[1]
    elif len(parts) == 3:
        t, r, b = parts[0], parts[1], parts[2]
        l = parts[1]
    else:
        t, r, b, l = parts[0], parts[1], parts[2], parts[3]
    return {"top": t, "right": r, "bottom": b, "left": l}


def _set_cell_margins(cell: _Cell, props: dict[str, str]) -> None:
    """Map CSS ``padding`` / ``padding-{side}`` to cell ``w:tcMar`` (D-TF-02).

    OOXML ``w:tcMar`` width unit is twentieths of a point (twips). Resolves
    each side to pt via ``_parse_size_to_pt`` then *20. Per-side overrides win
    over the shorthand. Missing/zero sides are skipped (Word inherits defaults).
    """
    sides: dict[str, str] = {}

    if "padding" in props:
        sides.update(_padding_shorthand_sides(props["padding"]))

    for side in _BORDER_SIDES:
        side_key = f"padding-{side}"
        if side_key in props:
            sides[side] = props[side_key]

    if not sides:
        return

    # Convert to twips; drop unparseable / non-positive.
    twips: dict[str, int] = {}
    for side, raw in sides.items():
        pt = _parse_size_to_pt(raw)
        if pt is None:
            continue
        w = int(round(pt * 20))
        if w > 0:
            twips[side] = w

    if not twips:
        return

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for side in _BORDER_SIDES:
        if side not in twips:
            continue
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(twips[side]))
        node.set(qn("w:type"), "dxa")


def _set_cell_vertical_alignment(cell: _Cell, alignment: str) -> None:
    """Set vertical alignment on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = tcPr.find(qn("w:vAlign"))
    if tcVAlign is None:
        tcVAlign = OxmlElement("w:vAlign")
        tcPr.append(tcVAlign)
    tcVAlign.set(qn("w:val"), alignment)


def _set_cell_borders(cell: _Cell, props: dict[str, str]) -> None:
    """Set borders on a single table cell from CSS properties."""
    border_info: dict[str, dict[str, str]] = {}

    if "border" in props:
        parsed = _parse_border_shorthand(props["border"])
        for side in _BORDER_SIDES:
            border_info[side] = dict(parsed)

    if "border-color" in props:
        for side in _BORDER_SIDES:
            border_info.setdefault(side, {})["color"] = props["border-color"]

    for side in _BORDER_SIDES:
        side_key = f"border-{side}"
        if side_key in props:
            parsed = _parse_border_shorthand(props[side_key])
            border_info.setdefault(side, {}).update(parsed)
        if f"border-{side}-color" in props:
            border_info.setdefault(side, {})["color"] = props[f"border-{side}-color"]

    if not border_info:
        return

    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for side in _BORDER_SIDES:
        info = border_info.get(side)
        if not info:
            continue
        # Build via shared border utility (preserves legacy cell color 000000).
        elem = build_border_element(side, info, default_color="000000")
        existing = tcBorders.find(qn(f"w:{side}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(elem)


def _set_cell_paragraph_format(cell: _Cell, props: dict[str, str], default_line_height: str = "1.5") -> None:
    """Set paragraph format for all paragraphs in a cell.

    Defaults to single line spacing to ensure consistent rendering
    with HTML layout behavior.

    If line-height and font-size are provided, calculates vertical padding as:
        before = after = (line-height - 1) / 2 * font-size
    This mirrors how browsers center text vertically within a line box.

    Args:
        default_line_height: Fallback line-height value (inherited from <body>)
            when the cell does not specify its own line-height.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Calculate before/after spacing from line-height and font-size
    before_twips: int | None = None
    after_twips: int | None = None
    # 如果没有line-height, 则取body中继承来的line-height作为默认值
    line_height_raw = props.get("line-height", default_line_height)
    font_size_raw = props.get("font-size", "12pt")

    if line_height_raw and font_size_raw:
        font_size_pt = _parse_size_to_pt(font_size_raw)
        if not font_size_pt:
            font_size_pt = 12
        # line-height can be a unitless multiplier (e.g. "1.5") or a length (e.g. "24px")
        lh_value = line_height_raw.strip()
        if re.match(r"^[\d.]+$", lh_value):
            # Unitless multiplier: line-height = multiplier * font-size
            line_height_pt = float(lh_value) * font_size_pt
        else:
            line_height_pt = _parse_size_to_pt(lh_value)

        if line_height_pt and line_height_pt > font_size_pt:
            # Distribute the extra space equally above and below
            half_leading_pt = (line_height_pt - font_size_pt) / 2
            # OOXML w:before / w:after unit: twentieths of a point (twips)
            before_twips = int(half_leading_pt * 20)
            after_twips = int(half_leading_pt * 20)

    for para in cell.paragraphs:
        pPr = para._p.get_or_add_pPr()
        # Set single line spacing: w:line=240 is a ratio factor (240/240 = 1×),
        # actual line height is determined by font size at render time.
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        # Apply before/after spacing derived from line-height, or clear to 0
        spacing.set(qn("w:before"), str(before_twips) if before_twips is not None else "0")
        spacing.set(qn("w:after"), str(after_twips) if after_twips is not None else "0")

        # text-indent → first_line_indent (overrides Normal style inheritance)
        if "text-indent" in props:
            from docx.shared import Cm
            indent_val = props["text-indent"].strip()
            if indent_val == "0" or indent_val == "0px" or indent_val == "0pt":
                # Explicitly clear indent inherited from Normal style
                ind = pPr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    pPr.append(ind)
                ind.set(qn("w:firstLine"), "0")
            else:
                from .style_mapper import _parse_size_to_cm
                cm = _parse_size_to_cm(indent_val)
                if cm is not None:
                    para.paragraph_format.first_line_indent = Cm(cm)

        # text-align → paragraph alignment
        if "text-align" in props:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            alignment = align_map.get(props["text-align"].strip().lower())
            if alignment is not None:
                para.alignment = alignment


def _apply_cell_text_styles(cell: _Cell, props: dict[str, str]) -> None:
    """Apply font-size, color, font-weight etc. to all runs in a cell."""
    from .style_mapper import apply_run_styles

    # Build a style string from relevant properties
    text_props = {}
    for key in ("font-size", "color", "font-family", "font-weight", "font-style", "text-decoration"):
        if key in props:
            text_props[key] = props[key]

    if not text_props:
        return

    style_str = "; ".join(f"{k}: {v}" for k, v in text_props.items())

    for para in cell.paragraphs:
        for run in para.runs:
            apply_run_styles(run, style_str)


# ---------------------------------------------------------------------------
# Row/Cell style helpers
# ---------------------------------------------------------------------------

def _get_row_border_props(html_row: Tag) -> dict[str, str]:
    """Extract border-related CSS props from a <tr> element's inline style.

    OOXML has no row-level border concept; these props will be distributed
    to each cell in the row as tcBorders.
    """
    row_style = html_row.get("style", "")
    if not row_style:
        return {}
    row_props = _parse_style(row_style)
    return {k: v for k, v in row_props.items() if k.startswith("border")}


def _build_cell_props(html_cell: Tag, row_border_props: dict[str, str]) -> dict[str, str]:
    """Build merged style props for a cell: cell's own style + inherited row border.

    Cell's own border declarations take priority over row-level border.
    """
    cell_style = html_cell.get("style", "")
    cell_props = _parse_style(cell_style) if cell_style else {}

    if row_border_props:
        merged_border = dict(row_border_props)
        # Cell's own border overrides row-level
        merged_border.update({k: v for k, v in cell_props.items() if k.startswith("border")})
        cell_props.update(merged_border)

    return cell_props


def _apply_cell_styles(docx_cell: "_Cell", cell_props: dict[str, str], default_line_height: str) -> None:
    """Apply all visual styles to a single docx table cell."""
    # Background color → cell shading
    bg_color = cell_props.get("background-color") or cell_props.get("background")
    if bg_color:
        color_hex = _parse_color(bg_color)
        if color_hex:
            _set_cell_shading(docx_cell, color_hex)

    # Vertical alignment (default center to align with HTML behavior)
    _set_cell_vertical_alignment(docx_cell, "center")

    # Cell borders
    has_border = any(k.startswith("border") for k in cell_props)
    if has_border:
        _set_cell_borders(docx_cell, cell_props)

    # Cell padding → w:tcMar (D-TF-02 缺口补遗)
    _set_cell_margins(docx_cell, cell_props)

    # Paragraph format (line spacing)
    _set_cell_paragraph_format(docx_cell, cell_props, default_line_height)

    # Text styles (font-size, color, etc.)
    _apply_cell_text_styles(docx_cell, cell_props)


# ---------------------------------------------------------------------------
# Public API: apply_table_styles (Phase 2d)
# ---------------------------------------------------------------------------

def apply_table_styles(document: DocxDocument, html: str) -> None:
    """Apply table/cell styles from HTML inline styles to docx tables.

    Matches docx tables 1:1 with HTML <table> elements (by document order).
    For each table:
      - Reads <table style="..."> → sets tblBorders
      - Reads <td>/<th> style="..."> → sets cell shading, borders, text styles
    """
    soup = BeautifulSoup(html, "lxml")
    html_tables = soup.find_all("table")
    docx_tables = document.tables

    # 传递默认行高
    default_line_height = "1.5"  # browser default fallback
    body_tag = soup.find("body")
    if body_tag:
        body_style = body_tag.get("style", "")
        if body_style:
            body_props = _parse_style(body_style)
            if "line-height" in body_props:
                default_line_height = body_props["line-height"]

    # Match by position (same order in both)
    for idx, docx_table in enumerate(docx_tables):
        if idx >= len(html_tables):
            break

        html_table: Tag = html_tables[idx]

        # --- Table-level styles (borders) ---
        table_style = html_table.get("style", "")
        if table_style:
            table_props = _parse_style(table_style)
            _set_table_borders(docx_table, table_props)

        # --- Table width: force 100% when CSS width is 100% or not set ---
        # html4docx often produces tblW=0/auto with tiny gridCol widths.
        # For document-targeted HTML (not web), tables should fill the
        # available page width by default.
        _force_table_full_width(docx_table, html_table)

        # --- Cell-level styles ---
        html_rows = html_table.find_all("tr")

        for row_idx, docx_row in enumerate(docx_table.rows):
            if row_idx >= len(html_rows):
                break

            html_row: Tag = html_rows[row_idx]
            row_border_props = _get_row_border_props(html_row)
            html_cells = html_row.find_all(["td", "th"])

            for col_idx, docx_cell in enumerate(docx_row.cells):
                if col_idx >= len(html_cells):
                    break

                html_cell: Tag = html_cells[col_idx]
                cell_props = _build_cell_props(html_cell, row_border_props)
                _apply_cell_styles(docx_cell, cell_props, default_line_height)

        # --- Remove leading empty paragraphs in cells (html4docx artifact) ---
        # html4docx turns whitespace between <td> and the first <p> into an
        # empty paragraph. Remove it if the cell has >1 paragraph and the
        # first one is empty.
        _strip_cell_leading_empty_paras(docx_table)
