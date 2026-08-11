"""Tests for callout rich-text preservation (US3 / FR-005 / D-R2).

Verifies ``render_callout`` no longer flattens content via ``get_text`` but
delegates to ``render_children_into_cell``, keeping a bold title paragraph,
multiple ``<p>`` as independent paragraphs, ``<li>`` items, and not corrupting
variant colors / emoji. Nested ``<table>`` must not crash.
"""
from __future__ import annotations

import warnings

from bs4 import BeautifulSoup
from docx import Document

from html_to_docx.components import render_component


def _el(html: str):
    soup = BeautifulSoup(html, "lxml")
    return soup.find(attrs={"data-component": True})


def _cell_paragraphs(doc):
    assert len(doc.tables) == 1, "callout should render exactly one 1x1 table"
    cell = doc.tables[0].cell(0, 0)
    # Non-empty paragraphs (independent blocks)
    return [p for p in cell.paragraphs if p.text.strip()], cell


# ---------------------------------------------------------------------------
# AC1: <strong> title + two <p> → >=3 independent paragraphs, first bold, no \n
# ---------------------------------------------------------------------------

def test_ac1_bold_title_and_multi_paragraphs():
    doc = Document()
    el = _el(
        '<div data-component="callout" data-variant="info">'
        '<div class="callout-content">'
        '<p><strong>核心结论</strong></p>'
        '<p>第一段说明文字，包含 <strong>加粗</strong> 行内强调。</p>'
        '<p>第二段补充文字，独立成段不应被拍平。</p>'
        '</div></div>'
    )
    assert render_component(el, doc) is True

    nonempty, cell = _cell_paragraphs(doc)
    # >= 3 independent paragraphs (title + 2 body)
    assert len(nonempty) >= 3, f"expected >=3 paragraphs, got {len(nonempty)}"

    # First paragraph must contain a bold run
    first = nonempty[0]
    assert any(r.font.bold is True for r in first.runs), "title paragraph must have a bold run"
    assert "核心结论" in first.text

    # Text must NOT be flattened with newline soft-wraps inside a single paragraph
    for p in nonempty:
        assert "\n" not in p.text, f"paragraph text should not contain flattened newline: {p.text!r}"

    # Inline <strong> inside body paragraph → a bold run exists somewhere
    all_runs = [r for p in nonempty for r in p.runs]
    assert any(r.font.bold and "加粗" in r.text for r in all_runs), "inline strong should produce bold run"


# ---------------------------------------------------------------------------
# AC2: <ul><li> → each <li> an independent paragraph
# ---------------------------------------------------------------------------

def test_ac2_list_items_independent_paragraphs():
    doc = Document()
    el = _el(
        '<div data-component="callout">'
        '<div class="callout-content">'
        '<ul><li>列表项一</li><li>列表项二</li><li>列表项三</li></ul>'
        '</div></div>'
    )
    assert render_component(el, doc) is True

    nonempty, _ = _cell_paragraphs(doc)
    texts = [p.text.strip() for p in nonempty]
    assert "列表项一" in texts
    assert "列表项二" in texts
    assert "列表项三" in texts
    # Each li is its own paragraph (>=3), not flattened into one
    assert len(nonempty) >= 3, f"expected >=3 li paragraphs, got {texts}"
    for t in texts:
        assert "列表项一列表项二" not in t, "list items must not be flattened together"


# ---------------------------------------------------------------------------
# AC3: variant color + emoji preserved (no regression)
# ---------------------------------------------------------------------------

def test_ac3_variant_color_and_emoji_preserved():
    doc = Document()
    el = _el(
        '<div data-component="callout" data-variant="warning">'
        '<div class="callout-content">'
        '<p>⚠️ 警示型 callout，保留 variant 颜色与 emoji。</p>'
        '</div></div>'
    )
    assert render_component(el, doc) is True

    cell = doc.tables[0].cell(0, 0)
    tc_xml = cell._tc.xml
    # warning variant border-left color E65100 and background fill FDF3EF preserved
    assert "E65100" in tc_xml, "warning border-left color must be preserved"
    assert "FDF3EF" in tc_xml.upper() or "fdf3ef" in tc_xml, "warning background shd must be preserved"
    # emoji preserved in text
    full_text = "".join(p.text for p in cell.paragraphs)
    assert "⚠️" in full_text, "emoji must be preserved, not stripped"
    assert "警示型" in full_text


# ---------------------------------------------------------------------------
# AC3b: inline style overrides variant default (regression guard)
# ---------------------------------------------------------------------------

def test_ac3b_inline_style_overrides_variant():
    doc = Document()
    el = _el(
        '<div data-component="callout" data-variant="info" '
        'style="background-color:#eef4fb;border-left-color:#1a4d8f">'
        '<div class="callout-content"><p>覆盖默认色</p></div></div>'
    )
    assert render_component(el, doc) is True
    tc_xml = doc.tables[0].cell(0, 0)._tc.xml.upper()
    assert "EEF4FB" in tc_xml, "inline background-color should override variant"
    assert "1A4D8F" in tc_xml, "inline border-left-color should override variant"


# ---------------------------------------------------------------------------
# Edge: nested <table> inside callout must not crash
# ---------------------------------------------------------------------------

def test_edge_nested_table_does_not_crash():
    doc = Document()
    el = _el(
        '<div data-component="callout">'
        '<div class="callout-content">'
        '<p>前置说明</p>'
        '<table><tr><td>单元格A</td><td>单元格B</td></tr></table>'
        '</div></div>'
    )
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        assert render_component(el, doc) is True  # must not raise
    cell = doc.tables[0].cell(0, 0)
    full_text = "".join(p.text for p in cell.paragraphs)
    assert "前置说明" in full_text


# ---------------------------------------------------------------------------
# Edge: callout with bare text content (legacy simple form) still works
# ---------------------------------------------------------------------------

def test_edge_bare_text_content():
    doc = Document()
    el = _el('<div data-component="callout"><div class="callout-content">纯文本内容</div></div>')
    assert render_component(el, doc) is True
    cell = doc.tables[0].cell(0, 0)
    full_text = "".join(p.text for p in cell.paragraphs)
    assert "纯文本内容" in full_text
