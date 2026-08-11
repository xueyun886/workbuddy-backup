"""test_paragraph_border_shading.py — US1: paragraph border / shading (FR-002/003 / SC-001).

Covers:
  - AC1: border-left:3px solid #1a4d8f → pPr/pBdr/left, w:sz=18 (3px*0.75*8≈18), w:color=1A4D8F
  - AC3: background-color:#eef4fb → pPr/shd @w:fill=EEF4FB
  - AC5: border:1px solid #ccc → pBdr top/bottom/left/right (four sides)
  - Edge: border-left:2px solid (no color) → w:color=auto, no error
  - Edge: border-left:2px dashed #999 → w:val=dashed
  - Edge: unknown style → single + warning
  - Edge: background:linear-gradient(...) → ignored + warning, no raise
  - FR-004 decision-lock handled in test_em_resolve / here (consecutive margins)
"""
from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from html_to_docx.style_mapper import apply_paragraph_styles


def _make_para(text: str = "test"):
    doc = Document()
    return doc.add_paragraph(text)


def _pBdr(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:pBdr"))


def _shd(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:shd"))


# --- AC1: single-side border with color ------------------------------------

def test_ac1_border_left_solid_color():
    para = _make_para()
    apply_paragraph_styles(para, "border-left:3px solid #1a4d8f")
    pBdr = _pBdr(para)
    assert pBdr is not None
    left = pBdr.find(qn("w:left"))
    assert left is not None
    assert left.get(qn("w:val")) == "single"
    # 3px → 2.25pt → *8 = 18
    assert left.get(qn("w:sz")) == "18"
    assert left.get(qn("w:color")) == "1A4D8F"
    # only left present (no other sides)
    assert pBdr.find(qn("w:right")) is None
    assert pBdr.find(qn("w:top")) is None


# --- AC3: background-color → shd --------------------------------------------

def test_ac3_background_color_shd():
    para = _make_para()
    apply_paragraph_styles(para, "background-color:#eef4fb")
    shd = _shd(para)
    assert shd is not None
    assert shd.get(qn("w:fill")) == "EEF4FB"
    assert shd.get(qn("w:val")) == "clear"
    assert shd.get(qn("w:color")) == "auto"


def test_background_shorthand_solid_color_shd():
    para = _make_para()
    apply_paragraph_styles(para, "background:#eef4fb")
    shd = _shd(para)
    assert shd is not None
    assert shd.get(qn("w:fill")) == "EEF4FB"


# --- AC5: border shorthand → four sides -------------------------------------

def test_ac5_border_shorthand_four_sides():
    para = _make_para()
    apply_paragraph_styles(para, "border:1px solid #ccc")
    pBdr = _pBdr(para)
    assert pBdr is not None
    for side in ("top", "bottom", "left", "right"):
        elem = pBdr.find(qn(f"w:{side}"))
        assert elem is not None, f"missing border side {side}"
        assert elem.get(qn("w:val")) == "single"
        assert elem.get(qn("w:color")) == "CCCCCC"


# --- Edge: missing color → auto, no raise -----------------------------------

def test_edge_border_missing_color_auto():
    para = _make_para()
    apply_paragraph_styles(para, "border-left:2px solid")
    pBdr = _pBdr(para)
    assert pBdr is not None
    left = pBdr.find(qn("w:left"))
    assert left is not None
    assert left.get(qn("w:color")) in ("auto", "000000")


# --- Edge: dashed style -----------------------------------------------------

def test_edge_border_dashed():
    para = _make_para()
    apply_paragraph_styles(para, "border-left:2px dashed #999")
    pBdr = _pBdr(para)
    left = pBdr.find(qn("w:left"))
    assert left.get(qn("w:val")) == "dashed"


# --- Edge: unknown style → single + warning ---------------------------------

def test_edge_unknown_style_falls_back_single():
    para = _make_para()
    with pytest.warns(UserWarning):
        apply_paragraph_styles(para, "border-left:2px groove #999")
    pBdr = _pBdr(para)
    left = pBdr.find(qn("w:left"))
    assert left.get(qn("w:val")) == "single"


# --- Edge: gradient background → ignored + warning, no raise ----------------

def test_edge_gradient_background_ignored_warns():
    para = _make_para()
    with pytest.warns(UserWarning):
        apply_paragraph_styles(para, "background:linear-gradient(90deg, #fff, #000)")
    # no shd written for gradient
    assert _shd(para) is None


def test_edge_url_background_ignored_no_raise():
    para = _make_para()
    with pytest.warns(UserWarning):
        apply_paragraph_styles(para, "background:url(bg.png)")
    assert _shd(para) is None


def test_edge_keyword_background_ignored():
    para = _make_para()
    with pytest.warns(UserWarning):
        apply_paragraph_styles(para, "background-color:transparent")
    assert _shd(para) is None


# --- FR-004 decision-lock: consecutive margins do NOT collapse --------------

def test_fr004_consecutive_margins_no_collapse():
    from docx.shared import Pt
    doc = Document()
    p1 = doc.add_paragraph("p1")
    p2 = doc.add_paragraph("p2")
    apply_paragraph_styles(p1, "margin-bottom:12pt")
    apply_paragraph_styles(p2, "margin-bottom:12pt")
    # Both faithfully map to space_after=12pt — no smart collapse/suppression.
    assert p1.paragraph_format.space_after == Pt(12)
    assert p2.paragraph_format.space_after == Pt(12)
