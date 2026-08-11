"""Section marker component renderer.

Renders ``data-component="section-marker"`` as a bold paragraph with
top margin and a thin bottom border — visually separates document sections.
"""
from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from bs4 import Tag

from . import register
from ..style_mapper import apply_run_styles, apply_paragraph_styles


@register("section-marker")
def render_section_marker(element: Tag, document: DocxDocument, anchor: Paragraph) -> None:
    """Render a section-marker component into *document*."""
    text = element.get_text(strip=True)
    number = element.get("data-number", "")
    level = element.get("data-level", "h2").lower()

    # Level → font size mapping (aligned with section-marker.html)
    _LEVEL_SIZE = {"h2": 18, "h3": 14, "h4": 12}
    font_size = _LEVEL_SIZE.get(level, 18)

    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(24)  # margin-top: ~2em
    para.paragraph_format.space_after = Pt(10)   # margin-bottom: ~0.8em

    # Number badge prefix
    if number:
        run_num = para.add_run(f"{number}  ")
        run_num.font.size = Pt(12)
        run_num.font.bold = True

    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(font_size)

    # Apply inline style from element
    style_attr = element.get("style", "")
    if style_attr:
        apply_paragraph_styles(para, style_attr)
        apply_run_styles(run, style_attr)

    # Thin bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    if not anchor is None:
        new_element = para._element
        new_element.getparent().remove(new_element)
        anchor._element.addprevious(new_element)
