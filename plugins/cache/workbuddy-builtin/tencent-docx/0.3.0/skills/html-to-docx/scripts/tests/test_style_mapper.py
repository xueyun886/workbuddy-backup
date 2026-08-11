"""Tests for html_to_docx.style_mapper."""
from __future__ import annotations

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from html_to_docx.style_mapper import (
    apply_font_family,
    apply_paragraph_styles,
    apply_run_styles,
)


def _make_run(text: str = "test"):
    doc = Document()
    para = doc.add_paragraph()
    return para.add_run(text)


def _make_para(text: str = "test"):
    doc = Document()
    return doc.add_paragraph(text)


def test_font_family_eastasia():
    """apply_font_family sets w:eastAsia on the run's rFonts element."""
    run = _make_run()
    apply_font_family(run, "仿宋_GB2312")
    rFonts = run._r.find(f".//{qn('w:rFonts')}")
    assert rFonts is not None
    assert rFonts.get(qn("w:eastAsia")) == "仿宋_GB2312"


def test_text_indent():
    """apply_paragraph_styles sets first_line_indent from text-indent."""
    para = _make_para()
    apply_paragraph_styles(para, "text-indent: 1cm")
    assert abs(para.paragraph_format.first_line_indent - Cm(1)) < 100  # ±100 EMU tolerance


def test_line_height():
    """apply_paragraph_styles sets line_spacing from line-height (pt)."""
    para = _make_para()
    apply_paragraph_styles(para, "line-height: 24pt")
    assert para.paragraph_format.line_spacing == Pt(24)


def test_font_size():
    """apply_run_styles sets font size from font-size."""
    run = _make_run()
    apply_run_styles(run, "font-size: 16pt")
    assert run.font.size == Pt(16)


def test_color():
    """apply_run_styles sets font color from color (#hex)."""
    run = _make_run()
    apply_run_styles(run, "color: #ff0000")
    assert run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_text_align():
    """apply_paragraph_styles sets paragraph alignment from text-align."""
    para = _make_para()
    apply_paragraph_styles(para, "text-align: center")
    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
