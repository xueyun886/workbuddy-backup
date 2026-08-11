"""header_footer.py — Render document cover area and OOXML footer.

Semantic mapping:
  <header>（无 class）        → 文档封面区，渲染为正文最前面的普通段落（只出现一次）
                                 包含：.doc-title / h1、.doc-subtitle、.doc-meta
  class="doc-footer"          → OOXML Section Footer，每页底部自动重复
                                 通过 section.footer 写入

TODO: class="doc-header"     → OOXML Section Header，每页顶部自动重复
                                 通过 section.header 写入（待实现）
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def render_header(soup_header, document: Document) -> None:
    """Insert header content at the *beginning* of *document* body.

    Reliable strategy:
      1. Append title / subtitle / meta via python-docx API, capturing
         each returned element's ._element reference directly.
      2. Move those elements to position 0 using body.insert(0, ...)
         in reverse order so the final sequence is title → subtitle → meta.

    Note: we capture ._element references IMMEDIATELY after each call,
    because python-docx may shift sectPr around and body position slicing
    is unreliable.
    """
    if soup_header is None:
        return

    body = document.element.body
    added: list = []  # lxml elements in insertion order

    title_el = soup_header.find(class_="doc-title") or soup_header.find("h1")
    if title_el:
        h = document.add_heading(title_el.get_text(strip=True), level=1)
        added.append(h._element)

    subtitle_el = soup_header.find(class_="doc-subtitle")
    if subtitle_el:
        para = document.add_paragraph()
        run = para.add_run(subtitle_el.get_text(strip=True))
        run.font.bold = True
        run.font.size = Pt(14)
        added.append(para._element)

    meta_el = soup_header.find(class_="doc-meta")
    if meta_el:
        para = document.add_paragraph()
        run = para.add_run(meta_el.get_text(strip=True))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        added.append(para._element)

    if not added:
        return

    # Move to front: insert in reverse so final order is title→subtitle→meta.
    # body.insert(0, elem) keeps sectPr at the end (lxml appends sectPr
    # after the inserted element automatically when needed).
    for elem in reversed(added):
        if elem.getparent() is body:
            body.remove(elem)
        body.insert(0, elem)


def render_footer(soup_footer, document: Document) -> None:
    """Set the document section footer from *soup_footer* element."""
    if soup_footer is None:
        return

    text = soup_footer.get_text(separator=" ").strip()
    section = document.sections[0]
    footer = section.footer
    # Use the existing footer paragraph (always present)
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = text
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Style the run
    if para.runs:
        run = para.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Add top border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "AAAAAA")
    pBdr.append(top)
    pPr.append(pBdr)
