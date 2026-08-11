"""element_mapper.py — Patch layer for HTML elements not covered by html-for-docx.

Covers: <caption>, <sup>/<sub>, colspan/rowspan merged cells, <hr>, <blockquote>.
"""
from __future__ import annotations

import warnings
from typing import TYPE_CHECKING

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from docx.table import Table, _Cell


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def patch_elements(document: Document, html: str) -> None:
    """Apply element patches to *document* using parsed *html* for context.

    Currently used for elements that html-for-docx does not handle or handles
    insufficiently. The main Converter calls this after the base conversion.
    """
    # Individual patch helpers are called from Converter as needed.
    # This function is kept as a no-op aggregator; callers use the helpers
    # directly for finer control.
    pass


# ---------------------------------------------------------------------------
# <caption> — rendered as a centred, muted paragraph above the table
# ---------------------------------------------------------------------------

def add_caption(document: Document, caption_text: str) -> None:
    """Insert a table caption paragraph at the end of the document."""
    para = document.add_paragraph(caption_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# <sup> / <sub>
# ---------------------------------------------------------------------------

def apply_superscript(run) -> None:
    """Mark *run* as superscript."""
    run.font.superscript = True


def apply_subscript(run) -> None:
    """Mark *run* as subscript."""
    run.font.subscript = True


# ---------------------------------------------------------------------------
# Table cell merging — colspan / rowspan
# ---------------------------------------------------------------------------

def merge_cells(table, row_idx: int, col_idx: int,
                rowspan: int = 1, colspan: int = 1) -> None:
    """Merge a rectangular block of cells starting at (row_idx, col_idx).

    Silently skips if indices are out of range.
    """
    try:
        start_cell = table.cell(row_idx, col_idx)
        end_cell   = table.cell(row_idx + rowspan - 1, col_idx + colspan - 1)
        start_cell.merge(end_cell)
    except IndexError:
        warnings.warn(
            f"merge_cells: out of range ({row_idx},{col_idx}) "
            f"+{rowspan}r +{colspan}c in {len(table.rows)}×{len(table.columns)} table",
            stacklevel=2,
        )


def apply_table_spans(document: Document, soup_table: Tag, doc_table) -> None:
    """Apply colspan/rowspan from a BeautifulSoup <table> tag to a docx Table.

    This is a best-effort implementation; complex nested spans may not be
    perfectly reproduced.
    """
    rows = soup_table.find_all("tr")
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells):
            rowspan = int(cell.get("rowspan", 1))
            colspan = int(cell.get("colspan", 1))
            if rowspan > 1 or colspan > 1:
                merge_cells(doc_table, r_idx, c_idx, rowspan, colspan)


# ---------------------------------------------------------------------------
# <hr> — horizontal rule as a paragraph with a bottom border
# ---------------------------------------------------------------------------

def add_horizontal_rule(document: Document) -> None:
    """Add a paragraph that visually represents <hr> via a bottom border."""
    para = document.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# <blockquote> — left-indented paragraph with a left border
# ---------------------------------------------------------------------------

def add_blockquote(document: Document, text: str) -> None:
    """Add a blockquote-styled paragraph."""
    from docx.shared import Cm
    para = document.add_paragraph(text)
    fmt = para.paragraph_format
    fmt.left_indent = Cm(1)

    # Add left border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "AAAAAA")
    pBdr.append(left)
    pPr.append(pBdr)
